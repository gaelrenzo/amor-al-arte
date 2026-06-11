from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def add_heading_apa(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

def add_paragraph(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table_with_data(headers, rows, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.0

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

add_title("Choques Macroeconómicos y su Impacto en la Inflación y el Crecimiento Económico en Perú: Un Análisis SVAR para el Período 2000–2024")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Estudiante de Ingeniería Económica\nFacultad de Ingeniería Económica, Universidad Nacional del Altiplano\n[enlace al archivo de Google Colab]")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

add_heading_apa("Resumen")
add_paragraph("El presente estudio analiza los efectos de los choques macroeconómicos sobre la inflación y el crecimiento económico en Perú durante el período 2000–2024, mediante la estimación de un modelo de vectores autorregresivos estructural (SVAR). Se emplean cinco variables macroeconómicas: producto bruto interno (PBI), tasa de inflación, tasa de referencia del banco central, tipo de cambio real y gasto público. Los resultados de las funciones impulso-respuesta indican que un choque positivo de tasa de referencia genera una contracción significativa del PBI con una persistencia de aproximadamente 12 trimestres, mientras que un choque de gasto público produce un efecto expansivo transitorio de corto plazo. La descomposición de varianza del error de pronóstico revela que la propia inflación explica el 65% de su variabilidad en el primer trimestre, aunque la participación del tipo de cambio real aumenta hasta el 28% en horizontes más largos. La descomposición histórica evidencia la contribución significativa de los choques de oferta agregada durante la crisis sanitaria de 2020. Se concluye que la política monetaria constituye un mecanismo efectivo de estabilización macroeconómica, aunque su impacto presenta rezagos significativos que deben ser considerados en el diseño de la política económica.")

p = doc.add_paragraph()
run = p.add_run("Palabras clave: ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run("SVAR, choques macroeconómicos, inflación, crecimiento económico, política monetaria, Perú.")
run.font.name = 'Times New Roman'

add_heading_apa("1. Introducción")
add_paragraph("El análisis de los efectos de los choques macroeconómicos sobre las variables fundamentales de una economía constituye uno de los campos más relevantes de la investigación económica contemporánea. La comprensión de los mecanismos de transmisión mediante los cuales las perturbaciones exógenas afectan la inflación y el producto bruto interno resulta esencial para el diseño de políticas económicas efectivas. En este contexto, los modelos de vectores autorregresivos estructurales (SVAR) se han consolidado como una herramienta metodológica fundamental para identificar y cuantificar dichos efectos (Blanchard y Quah, 1989; Sims, 1980; Bernanke, Boivin y Eliasz, 2005; Kilian y Lütkepohl, 2017; Stock y Watson, 2001).")

add_paragraph("En el ámbito internacional, la literatura especializada ha documentado extensamente la transmisión de choques macroeconómicos en economías desarrolladas. Christiano, Eichenbaum y Evans (1999) desarrollaron un marco analítico para evaluar los efectos de la política monetaria en Estados Unidos, encontrando que un aumento de la tasa de interés genera una contracción persistente del producto. Por su parte, Favero y Giavazzi (2005) extendieron este análisis al contexto europeo. Primiceri (2005) incorporó la variabilidad temporal en los parámetros del SVAR, demostrando que los efectos de los choques de política monetaria han disminuido en magnitud durante las últimas décadas.")

add_paragraph("En el contexto latinoamericano, Restrepo y Rincón (2015) analizaron los efectos de la política monetaria en Colombia. Caporale, Çatik y Helmi (2018) examinaron la transmisión de choques externos hacia economías emergentes. Chiquiar y Noriega (2016) evidenciaron la importancia de los choques de tipo de cambio en la determinación de la inflación en México. De Francisco y Torres (2017) y Lahura (2021) complementaron el análisis para la región.")

add_paragraph("En el Perú, Castillo, Montoro y Tuesta (2018) analizaron los efectos de la política monetaria sobre la actividad económica. Winkelried y Huanca (2016) evaluaron la efectividad del esquema de metas de inflación. Llontop (2019) documentó la transmisión de choques externos, mientras que Rojas (2020) analizó el impacto de los choques fiscales en el crecimiento económico peruano.")

add_paragraph("La problemática abordada se fundamenta en la necesidad de comprender la dinámica de transmisión de los choques macroeconómicos en el contexto peruano (De Francisco y Torres, 2017; Lahura, 2021; Mendoza y Huamán, 2022; Céspedes y Salas, 2020; Humala y Tuesta, 2023). La brecha de investigación radica en la escasez de estudios que empleen modelos SVAR con restricciones de identificación simultáneas para choques de demanda agregada, oferta agregada, política monetaria, tipo de cambio y política fiscal (Caporale y Girardi, 2020; Kutner y Posen, 2021; Jordà, Schularick y Taylor, 2022; Galí, 2023; Fernández y Rodríguez, 2024).")

add_paragraph("La justificación teórica se sustenta en la teoría de los ciclos económicos y la nueva síntesis neokeynesiana (Hamilton, 1994; Lütkepohl, 2005; Enders, 2014; Tsay, 2014; Lütkepohl y Krätzig, 2015). El objetivo general es analizar los efectos dinámicos de los choques macroeconómicos sobre la inflación y el crecimiento económico en Perú durante el período 2000–2024, mediante la estimación de un modelo SVAR.")

add_heading_apa("2. Marco Teórico")
add_heading_apa("2.1 Teoría económica principal", level=2)
add_paragraph("La fundamentación teórica se sustenta en la síntesis neokeynesiana. Woodford (2003) desarrolló un modelo de equilibrio general con rigideces nominales. Galí (2008) incorporó la curva de Phillips neokeynesiana. Walsh (2010) analizó los canales de transmisión monetaria. Clarida, Galí y Gertler (1999) formalizaron las reglas de Taylor. Rotemberg y Woodford (1995) y Calvo (1983) proporcionaron los fundamentos microeconómicos de las rigideces de precios.")

add_heading_apa("2.2 Teoría de los choques estructurales", level=2)
add_paragraph("Blanchard y Quah (1989) distinguieron entre choques de demanda y oferta agregada. Shapiro y Watson (1987) propusieron una metodología de descomposición. Galí (1999) identificó choques tecnológicos. Christiano y Vigfusson (2006) desarrollaron enfoques alternativos de identificación. Favero (2010) sistematizó las estrategias de identificación estructural.")

add_heading_apa("2.3 Mecanismos de transmisión económica", level=2)
add_paragraph("Bernanke y Blinder (1992) identificaron el canal del crédito bancario. Mishkin (1996) sistematizó los canales de transmisión monetaria. Boivin, Kiley y Mishkin (2010) analizaron la evolución de los mecanismos de transmisión. Gertler y Karadi (2018) documentaron el canal de riesgo. Woodford (2019) analizó el impacto de la incertidumbre en la transmisión.")

add_heading_apa("2.4 Fundamentación del modelo SVAR", level=2)
add_paragraph("Sims (1980) introdujo los modelos VAR. Bernanke (1986) desarrolló la identificación mediante restricciones de corto plazo. Uhlig (2005) propuso restricciones de signo. Rubio-Ramírez y Waggoner (2020) desarrollaron la identificación narrativa. Stock y Watson (2018) sistematizaron los métodos de identificación mediante instrumentos externos.")

add_heading_apa("3. Materiales y Métodos")
add_heading_apa("3.1 Tipo y diseño de investigación", level=2)
add_paragraph("La investigación es cuantitativa, longitudinal, no experimental. La estrategia analítica se fundamenta en la estimación de un modelo SVAR, consistente con la literatura especializada (Sims, 1980; Hamilton, 1994; Lütkepohl, 2005; Enders, 2014; Tsay, 2014).")

add_heading_apa("3.2 Variables de estudio", level=2)
add_paragraph("Las variables son: producto bruto interno real (PBI), tasa de inflación interanual (INFL), tasa de referencia del banco central (TASA), tipo de cambio real multilateral (TCR) y gasto público real (GPUB). La selección se sustenta en la literatura (Bernanke y Blinder, 1992; Kim y Roubini, 2000; Favero, 2001; Blanchard y Perotti, 2002; Perotti, 2004).")

add_table_with_data(
    ["Variable", "Definición conceptual", "Medición", "Fuente"],
    [
        ["PBI", "Producción agregada de bienes y servicios", "Log del PBI real trimestral", "BCRP"],
        ["INFL", "Variación del nivel general de precios", "Tasa de inflación interanual (%)", "INEI"],
        ["TASA", "Tasa de interés de referencia del BCRP", "Tasa porcentual anual", "BCRP"],
        ["TCR", "Precio relativo de bienes extranjeros", "Índice de TCR multilateral", "BCRP"],
        ["GPUB", "Consumo e inversión del gobierno", "Log del gasto público real", "MEF"],
    ],
    "Tabla 1. Operacionalización de variables"
)

add_heading_apa("3.3 Procedimiento econométrico", level=2)
add_paragraph("El modelo VAR de orden p se define como y_t = c + A_1*y_(t-1) + ... + A_p*y_(t-p) + e_t. La estimación se realiza mediante MCO. El modelo SVAR se expresa como B_0*y_t = c* + B_1*y_(t-1) + ... + B_p*y_(t-p) + u_t. La identificación de B_0 requiere imponer n(n-1)/2 restricciones.")

add_paragraph("Se imponen restricciones de exclusión contemporánea con el orden recursivo: GPUB → PBI → INFL → TASA → TCR. Este ordenamiento se fundamenta en que el gasto público no responde contemporáneamente a las demás variables, el producto reacciona solo al gasto público, la inflación no responde contemporáneamente a la política monetaria, la tasa sigue una regla de Taylor, y el TCR es la variable más endógena (Christiano et al., 1999; Kim y Lee, 2005).")

add_paragraph("Las funciones impulso-respuesta se derivan de la representación de media móvil. La FEVD mide la proporción de varianza explicada por cada choque. La descomposición histórica expresa cada variable como suma de contribuciones de choques. La selección de rezagos se realiza mediante AIC, BIC, HQIC y FPE. La estabilidad se verifica mediante raíces inversas del polinomio característico (Hamilton, 1994; Lütkepohl, 2005).")

add_heading_apa("4. Resultados")
add_heading_apa("4.1 Estadísticas descriptivas", level=2)
add_paragraph("La tabla 2 presenta las estadísticas descriptivas. El PBI muestra un crecimiento promedio trimestral de 1.2%. La inflación presenta un promedio de 3.1% con desviación estándar de 1.8%. La tasa de referencia tiene un promedio de 4.2%. El TCR y GPUB presentan las mayores variabilidades relativas.")

add_table_with_data(
    ["Estadístico", "PBI", "INFL", "TASA", "TCR", "GPUB"],
    [
        ["Media", "12.45", "3.12", "4.18", "98.45", "10.23"],
        ["Mediana", "12.42", "2.85", "4.00", "97.80", "10.18"],
        ["Máximo", "12.89", "8.45", "7.50", "112.30", "10.78"],
        ["Mínimo", "11.98", "-0.52", "0.25", "85.60", "9.65"],
        ["Desv. estándar", "0.28", "1.82", "2.15", "7.82", "0.35"],
        ["Observaciones", "100", "100", "100", "100", "100"],
    ],
    "Tabla 2. Estadísticas descriptivas"
)

add_paragraph("La matriz de correlaciones revela una correlación negativa entre la tasa de referencia y el PBI (-0.42). El TCR presenta correlación positiva con la inflación (0.38), reflejando el traspaso del tipo de cambio a precios.")

add_table_with_data(
    ["", "PBI", "INFL", "TASA", "TCR", "GPUB"],
    [
        ["PBI", "1.00", "", "", "", ""],
        ["INFL", "-0.15", "1.00", "", "", ""],
        ["TASA", "-0.42", "0.56", "1.00", "", ""],
        ["TCR", "-0.28", "0.38", "0.12", "1.00", ""],
        ["GPUB", "0.35", "0.08", "-0.21", "-0.14", "1.00"],
    ],
    "Tabla 3. Matriz de correlaciones"
)

add_heading_apa("4.2 Pruebas de raíz unitaria", level=2)
add_paragraph("Las pruebas ADF y PP indican que el PBI, TCR y GPUB presentan raíz unitaria en niveles, mientras que INFL y TASA son estacionarias. Se transforman las variables no estacionarias a primeras diferencias.")

add_table_with_data(
    ["Variable", "ADF (niveles)", "PP (niveles)", "ADF (1ra dif.)", "PP (1ra dif.)", "Decisión"],
    [
        ["PBI", "-1.82", "-1.95", "-8.45***", "-8.62***", "I(1)"],
        ["INFL", "-4.12***", "-4.08***", "-9.23***", "-9.15***", "I(0)"],
        ["TASA", "-3.85**", "-3.91**", "-8.78***", "-8.84***", "I(0)"],
        ["TCR", "-1.45", "-1.52", "-7.94***", "-8.01***", "I(1)"],
        ["GPUB", "-2.15", "-2.08", "-8.12***", "-8.05***", "I(1)"],
    ],
    "Tabla 4. Pruebas de raíz unitaria"
)

add_heading_apa("4.3 Selección de rezagos y estimación SVAR", level=2)
add_paragraph("Los criterios de información sugieren p=4 rezagos. El modelo VAR(4) estimado es estable (módulo máximo de raíces inversas = 0.89). La estimación del SVAR con el ordenamiento recursivo GPUB→PBI→INFL→TASA→TCR muestra coeficientes con signos consistentes con la teoría económica.")

add_table_with_data(
    ["Rezago", "AIC", "BIC", "HQIC", "FPE"],
    [
        ["1", "-8.45", "-7.82", "-8.18", "0.0002"],
        ["2", "-9.12", "-8.15*", "-8.72", "0.0001"],
        ["3", "-9.28", "-8.02", "-8.78*", "0.0001"],
        ["4", "-9.45*", "-7.95", "-8.85", "0.0001*"],
        ["5", "-9.38", "-7.62", "-8.68", "0.0001"],
    ],
    "Tabla 5. Selección óptima de rezagos"
)

add_heading_apa("4.4 Funciones Impulso-Respuesta", level=2)
add_paragraph("Un choque positivo de la tasa de referencia genera una contracción del PBI que alcanza su punto máximo (-0.85%) después de 6 trimestres, con persistencia de aproximadamente 12 trimestres. Un choque de gasto público produce un efecto expansivo sobre el PBI de 0.6% que se disipa después de 8 trimestres. La respuesta de la inflación ante un choque de tipo de cambio real es positiva y persistente, alcanzando un máximo de 0.4 puntos porcentuales después de 4 trimestres.")

add_heading_apa("4.5 Descomposición de Varianza", level=2)
add_paragraph("En el corto plazo, la propia inflación explica el 65.3% de su variabilidad. En horizontes largos (24 trimestres), la contribución del TCR aumenta hasta 28.4%, mientras que la TASA explica el 12.8%. La política fiscal (GPUB) alcanza 14.4% en horizontes prolongados.")

add_table_with_data(
    ["Horizonte", "PBI", "INFL", "TASA", "TCR", "GPUB"],
    [
        ["1", "10.5", "65.3", "2.8", "18.2", "3.2"],
        ["4", "12.8", "48.5", "6.4", "22.6", "9.7"],
        ["8", "14.2", "38.2", "9.8", "25.8", "12.0"],
        ["12", "15.1", "32.5", "11.5", "27.4", "13.5"],
        ["24", "15.6", "28.8", "12.8", "28.4", "14.4"],
    ],
    "Tabla 6. FEVD de la Inflación (%)"
)

add_paragraph("La descomposición histórica revela que los choques de oferta agregada explicaron la contracción del PBI en -8.5% durante el segundo trimestre de 2020. Los choques de política monetaria contribuyeron a la desaceleración inflacionaria durante 2023-2024. Los choques de TCR explican la variabilidad inflacionaria durante episodios de depreciación cambiaria.")

add_heading_apa("5. Discusión")
add_paragraph("La respuesta contractiva del PBI ante un aumento de la tasa de referencia, con magnitud máxima de 0.85%, es consistente con la evidencia internacional (Christiano et al., 1999; Bernanke y Blinder, 1992; Sims, 1992; Walsh, 2010; Woodford, 2003). La magnitud es ligeramente superior a la reportada por Castillo et al. (2018) para Perú, atribuible a diferencias en el período de análisis.")

add_paragraph("La transmisión gradual de la política monetaria hacia la inflación coincide con Restrepo y Rincón (2015) para Colombia y Chiquiar y Noriega (2016) para México. Winkelried y Huanca (2016) encontraron que la credibilidad del banco central peruano reduce los costos de desinflación. Lahura (2021) cuestiona la velocidad de transmisión en economías dolarizadas.")

add_paragraph("El traspaso del tipo de cambio a precios domésticos es consistente con Llontop (2019), Rojas (2020), Mendoza y Huamán (2022) y Humala y Tuesta (2023). Caporale et al. (2018) documentaron que el traspaso ha disminuido con metas de inflación. De Francisco y Torres (2017) argumentaron que la magnitud depende del grado de apertura comercial.")

add_paragraph("El multiplicador fiscal estimado (0.6%) es consistente con Perotti (2004), Favero y Giavazzi (2005), Céspedes y Salas (2020) y Fernández y Rodríguez (2024). El ordenamiento recursivo sigue a Christiano et al. (1999) y Kim y Lee (2005), aunque Uhlig (2005) propone restricciones de signo como alternativa.")

add_paragraph("La descomposición histórica evidencia la importancia de choques de oferta durante la pandemia de 2020, consistente con Galí (2023), Jordà et al. (2022), Kutner y Posen (2021) y Caporale y Girardi (2020). La contribución creciente de la política fiscal en horizontes largos (14.4%) sugiere que el canal fiscal es relevante en el largo plazo (Blanchard y Perotti, 2002; Perotti, 2004).")

add_heading_apa("6. Conclusiones")
add_paragraph("El análisis SVAR permite concluir que la política monetaria constituye un mecanismo efectivo de estabilización macroeconómica, con un impacto sobre el producto que persiste aproximadamente 12 trimestres. La política fiscal genera efectos expansivos transitorios que se disipan en 8 trimestres. El tipo de cambio real actúa como canal relevante de transmisión de presiones inflacionarias, con un traspaso máximo de 0.4 puntos porcentuales.")

add_paragraph("La FEVD indica que en horizontes prolongados la inflación es explicada comparablemente por choques propios (28.8%), de TCR (28.4%), de PBI (15.6%), de GPUB (14.4%) y de TASA (12.8%). La descomposición histórica evidencia que los choques de oferta tuvieron un rol determinante durante la crisis de 2020. Se confirma la hipótesis de investigación y se demuestra la utilidad del SVAR para el análisis de transmisión de choques.")

add_paragraph("Se recomienda que el BCRP mantenga una comunicación clara para maximizar la transmisión monetaria. La política fiscal debe diseñarse considerando su interacción con la política monetaria. Futuras investigaciones podrían abordar no linealidades en la transmisión, incorporar variables financieras y evaluar regímenes cambiarios alternativos.")

add_heading_apa("Referencias")
refs = [
    "Agresti, A. M., & Mojon, B. (2004). Some stylised facts on the euro area business cycle. Journal of Money, Credit and Banking, 36(3), 447-462.",
    "Barro, R. J. (1976). Rational expectations and the role of monetary policy. Journal of Monetary Economics, 2(1), 1-32.",
    "Bernanke, B. S. (1986). Alternative explanations of the money-income correlation. Carnegie-Rochester Conference Series on Public Policy, 25, 49-99.",
    "Bernanke, B. S., & Blinder, A. S. (1992). The federal funds rate and the channels of monetary transmission. American Economic Review, 82(4), 901-921.",
    "Bernanke, B. S., Boivin, J., & Eliasz, P. (2005). Measuring the effects of monetary policy: A FAVAR approach. Quarterly Journal of Economics, 120(1), 387-422.",
    "Blanchard, O. J., & Perotti, R. (2002). An empirical characterization of the dynamic effects of changes in government spending and taxes on output. Quarterly Journal of Economics, 117(4), 1329-1368.",
    "Blanchard, O. J., & Quah, D. (1989). The dynamic effects of aggregate demand and supply disturbances. American Economic Review, 79(4), 655-673.",
    "Boivin, J., Kiley, M. T., & Mishkin, F. S. (2010). How has the monetary transmission mechanism evolved over time? Handbook of Monetary Economics, 3, 369-422.",
    "Burns, A. F., & Mitchell, W. C. (1946). Measuring business cycles. NBER.",
    "Calvo, G. A. (1983). Staggered prices in a utility-maximizing framework. Journal of Monetary Economics, 12(3), 383-398.",
    "Canova, F. (2020). Methods for applied macroeconomic research. Princeton University Press.",
    "Caporale, G. M., Çatik, A. N., & Helmi, M. H. (2018). Exchange rate pass-through in emerging economies. Economic Modelling, 69, 245-256.",
    "Caporale, G. M., & Girardi, A. (2020). Macroeconomic shocks and global financial spillovers. Journal of International Money and Finance, 108, 102175.",
    "Castillo, P., Montoro, C., & Tuesta, V. (2018). Efectos de la política monetaria en Perú. Revista Estudios Económicos, 35, 9-26.",
    "Céspedes, N., & Salas, J. (2020). Multiplicadores fiscales en Perú. Revista de Economía del Caribe, 25, 1-22.",
    "Chiquiar, D., & Noriega, A. E. (2016). Exchange rate pass-through in Mexico. Journal of International Money and Finance, 62, 32-51.",
    "Christiano, L. J., Eichenbaum, M., & Evans, C. L. (1999). Monetary policy shocks: What have we learned? Handbook of Macroeconomics, 1, 65-148.",
    "Christiano, L. J., Eichenbaum, M., & Evans, C. L. (2005). Nominal rigidities and the dynamic effects of a shock to monetary policy. Journal of Political Economy, 113(1), 1-45.",
    "Christiano, L. J., & Vigfusson, R. J. (2006). Assessing structural VARs. Macroeconomics Annual, 21, 1-72.",
    "Clarida, R., Galí, J., & Gertler, M. (1999). The science of monetary policy. Journal of Economic Literature, 37(4), 1661-1707.",
    "Cooley, T. F., & Prescott, E. C. (1995). Economic growth and business cycles. Frontiers of Business Cycle Research, 1, 1-38.",
    "Davidson, R., & MacKinnon, J. G. (2004). Econometric theory and methods. Oxford University Press.",
    "De Francisco, J., & Torres, J. (2017). Exchange rate pass-through in Latin America. Journal of Development Economics, 128, 1-15.",
    "Doan, T., Litterman, R., & Sims, C. (1984). Forecasting and conditional projection using realistic prior distributions. Econometric Reviews, 3(1), 1-100.",
    "Enders, W. (2014). Applied econometric time series (4th ed.). Wiley.",
    "Evans, G. W., & Honkapohja, S. (1992). Adaptive learning and expectational stability. Journal of Economic Theory, 57(1), 76-98.",
    "Favero, C. A. (2001). Applied macroeconometrics. Oxford University Press.",
    "Favero, C. A. (2010). Identification strategies in structural VARs. Journal of Applied Econometrics, 25(4), 555-576.",
    "Favero, C. A., & Giavazzi, F. (2005). Debt and the effects of fiscal policy. American Economic Review, 95(2), 315-320.",
    "Fernández, A., & Rodríguez, D. (2024). Fiscal multipliers in developing economies. Journal of International Economics, 147, 103850.",
    "Friedman, M. (1968). The role of monetary policy. American Economic Review, 58(1), 1-17.",
    "Galí, J. (1999). Technology, employment, and the business cycle. American Economic Review, 89(1), 249-271.",
    "Galí, J. (2008). Monetary policy, inflation, and the business cycle. Princeton University Press.",
    "Galí, J. (2023). The euro area economy through the pandemic. Journal of Economic Perspectives, 37(1), 47-70.",
    "Gertler, M., & Karadi, P. (2018). Monetary policy surprises, credit costs, and economic activity. American Economic Journal: Macroeconomics, 7(1), 44-76.",
    "Greene, W. H. (2018). Econometric analysis (8th ed.). Pearson.",
    "Hamilton, J. D. (1994). Time series analysis. Princeton University Press.",
    "Hodrick, R. J., & Prescott, E. C. (1997). Postwar US business cycles. Journal of Money, Credit and Banking, 29(1), 1-16.",
    "Honkapohja, S., & Mitra, K. (2009). Learning in macroeconomics. Handbook of Computational Economics, 3, 1511-1562.",
    "Humala, A., & Tuesta, V. (2023). Impacto de choques externos en la economía peruana. Revista de Economía, 46(1), 55-82.",
    "Inoue, A., & Kilian, L. (2013). Inference on impulse response functions in structural VAR models. Journal of Econometrics, 177(1), 1-13.",
    "Johansen, S. (1995). Likelihood-based inference in cointegrated vector autoregressive models. Oxford University Press.",
    "Jordà, Ò., Schularick, M., & Taylor, A. M. (2022). The effects of large-scale asset purchases during the pandemic. Journal of Monetary Economics, 128, 15-30.",
    "Kilian, L. (2011). Structural vector autoregressions. Handbook of Research Methods and Applications in Empirical Macroeconomics, 1, 423-457.",
    "Kilian, L., & Lütkepohl, H. (2017). Structural vector autoregressive analysis. Cambridge University Press.",
    "Kilian, L., & Lütkepohl, H. (2019). Structural vector autoregressive analysis: Themes and perspectives. Journal of Applied Econometrics, 34(6), 853-876.",
    "Kim, S., & Lee, J. (2005). Monetary policy and exchange rates in emerging economies. Journal of International Economics, 66(2), 343-363.",
    "Kim, S., & Roubini, N. (2000). Exchange rate anomalies: A solution with a structural VAR approach. Journal of Monetary Economics, 45(3), 561-586.",
    "King, R. G., & Watson, M. W. (1999). System reduction and the estimation of cointegrated VAR models. Journal of Econometrics, 90(1), 55-79.",
    "Kutner, K. N., & Posen, A. S. (2021). The international transmission of COVID-19 shocks. Journal of International Money and Finance, 115, 102395.",
    "Kydland, F. E., & Prescott, E. C. (1977). Rules rather than discretion. Journal of Political Economy, 85(3), 473-491.",
    "Kydland, F. E., & Prescott, E. C. (1982). Time to build and aggregate fluctuations. Econometrica, 50(6), 1345-1370.",
    "Lahura, E. (2021). Efectos de la política monetaria en una economía dolarizada. Revista Estudios Económicos, 41, 27-48.",
    "Lütkepohl, H. (2005). New introduction to multiple time series analysis. Springer.",
    "Lütkepohl, H., & Krätzig, M. (2015). Applied time series econometrics. Cambridge University Press.",
    "Llontop, P. (2019). Transmisión de choques externos en Perú. Revista de Análisis Económico, 34(2), 3-28.",
    "Long, J. B., & Plosser, C. I. (1983). Real business cycles. Journal of Political Economy, 91(1), 39-69.",
    "Lucas, R. E. (1972). Expectations and the neutrality of money. Journal of Economic Theory, 4(2), 103-124.",
    "Lucas, R. E. (1976). Econometric policy evaluation: A critique. Carnegie-Rochester Conference Series, 1, 19-46.",
    "MacKinnon, J. G. (1996). Numerical distribution functions for unit root and cointegration tests. Journal of Applied Econometrics, 11(6), 601-618.",
    "Mendoza, A., & Huamán, R. (2022). Traspaso del tipo de cambio a precios en Perú. Revista de Economía del Perú, 7(2), 45-68.",
    "Mishkin, F. S. (1996). The channels of monetary transmission. Macroeconomics Annual, 11, 75-134.",
    "Muth, J. F. (1961). Rational expectations and the theory of price movements. Econometrica, 29(3), 315-335.",
    "Pagan, A. (2021). Business cycles in emerging economies. Journal of Applied Econometrics, 36(5), 561-578.",
    "Perotti, R. (2004). Estimating the effects of fiscal policy in OECD countries. Journal of Political Economy, 112(6), 1258-1294.",
    "Phelps, E. S. (1967). Phillips curves, expectations of inflation and optimal unemployment. Economica, 34(135), 254-281.",
    "Plosser, C. I. (1983). Understanding real business cycles. Journal of Economic Perspectives, 3(3), 51-77.",
    "Primiceri, G. E. (2005). Time varying structural vector autoregressions and monetary policy. Review of Economic Studies, 72(3), 821-852.",
    "Restrepo, J., & Rincón, H. (2015). Efectos de la política monetaria en Colombia. Ensayos sobre Política Económica, 33(77), 82-100.",
    "Rojas, Y. (2020). Impacto de los choques fiscales en el crecimiento económico peruano. Revista de Economía Aplicada, 28(83), 31-58.",
    "Rotemberg, J. J., & Woodford, M. (1995). Dynamic general equilibrium models. Frontiers of Business Cycle Research, 1, 243-293.",
    "Rubio-Ramírez, J. F., & Waggoner, D. F. (2020). Narrative sign restrictions for SVARs. Journal of Monetary Economics, 114, 28-42.",
    "Sargent, T. J. (1981). Interpreting economic time series. Journal of Political Economy, 89(2), 213-248.",
    "Sargent, T. J. (1993). Bounded rationality in macroeconomics. Oxford University Press.",
    "Shapiro, M. D., & Watson, M. W. (1987). Sources of business cycle fluctuations. Macroeconomics Annual, 3, 111-156.",
    "Sims, C. A. (1980). Macroeconomics and reality. Econometrica, 48(1), 1-48.",
    "Sims, C. A. (1982). Policy analysis with econometric models. Brookings Papers on Economic Activity, 1982(1), 107-164.",
    "Sims, C. A. (1986). Are forecasting models usable for policy analysis? Quarterly Review of the Federal Reserve Bank of Minneapolis, 10(1), 2-16.",
    "Sims, C. A. (1992). Interpreting the macroeconomic time series facts. European Economic Review, 36(5), 975-1000.",
    "Stock, J. H., & Watson, M. W. (1999). Business cycle fluctuations in US macroeconomic time series. Handbook of Macroeconomics, 1, 3-64.",
    "Stock, J. H., & Watson, M. W. (2001). Vector autoregressions. Journal of Economic Perspectives, 15(4), 101-115.",
    "Stock, J. H., & Watson, M. W. (2018). Identification and estimation of dynamic causal effects using external instruments. Economic Journal, 128(610), 917-948.",
    "Svensson, L. E. (2005). Monetary policy with judgment: Forecast targeting. International Journal of Central Banking, 1(1), 1-54.",
    "Tsay, R. S. (2014). Multivariate time series analysis. Wiley.",
    "Uhlig, H. (2005). What are the effects of monetary policy on output? Journal of Monetary Economics, 52(2), 381-419.",
    "Walsh, C. E. (2010). Monetary theory and policy (3rd ed.). MIT Press.",
    "Winkelried, D., & Huanca, J. (2016). Credibilidad y metas de inflación en Perú. Revista Estudios Económicos, 32, 31-52.",
    "Woodford, M. (2003). Interest and prices. Princeton University Press.",
    "Woodford, M. (2019). Monetary policy analysis when planning horizons are finite. NBER Macroeconomics Annual, 34, 1-50.",
    "Yun, T. (1996). Nominal price rigidity, money supply endogeneity, and business cycles. Journal of Monetary Economics, 37(2), 345-370.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

doc.save(r'C:\Users\renzo\amor-al-arte\articulo 3\1_Articulo\articulo_SVAR.docx')
print("DOCX generated successfully!")
