import os
import sys

def extract_pdf(pdf_path, txt_path):
    print(f"Extracting PDF: {pdf_path}")
    try:
        import pypdf
    except ImportError:
        print("pypdf is not installed. Installing pypdf...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pypdf"])
        import pypdf
        
    reader = pypdf.PdfReader(pdf_path)
    text = ""
    for i, page in enumerate(reader.pages):
        text += f"--- Page {i+1} ---\n"
        text += page.extract_text() or ""
        text += "\n"
        
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"Saved to {txt_path}")

def extract_docx(docx_path, txt_path):
    print(f"Extracting DOCX: {docx_path}")
    try:
        import docx
    except ImportError:
        print("python-docx is not installed. Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        
    doc = docx.Document(docx_path)
    text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        text.append(para.text)
        
    # Read tables
    for i, table in enumerate(doc.tables):
        text.append(f"\n--- Table {i+1} ---")
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            text.append(" | ".join(row_text))
            
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print(f"Saved to {txt_path}")

if __name__ == "__main__":
    pdf_file = "Pautas practica 10.pdf"
    docx_file = "Recomendaciones y rúbrica.docx"
    
    if os.path.exists(pdf_file):
        extract_pdf(pdf_file, "pautas_practica_10.txt")
    else:
        print(f"File not found: {pdf_file}")
        
    if os.path.exists(docx_file):
        extract_docx(docx_file, "recomendaciones_rubrica.txt")
    else:
        # Check if file has variations in spelling
        found = False
        for f in os.listdir("."):
            if "recomendaciones" in f.lower() and f.endswith(".docx"):
                extract_docx(f, "recomendaciones_rubrica.txt")
                found = True
                break
        if not found:
            print(f"File not found: {docx_file}")
