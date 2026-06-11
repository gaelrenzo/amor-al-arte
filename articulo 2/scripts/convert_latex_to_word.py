import os
import re
import shutil
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "D3D3D3", "space": "0"},
        bottom={"sz": 12, "color": "00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def clean_latex(text):
    if not text:
        return ""
    # Replace common LaTeX characters and markup
    text = text.replace(r'\%', '%')
    text = text.replace(r'\$', '$')
    text = text.replace(r'\&', '&')
    text = text.replace(r'\_', '_')
    text = text.replace(r'~', ' ')
    text = text.replace(r'\deg', '°')
    text = text.replace(r'\pm', '±')
    text = text.replace(r'\infty', '∞')
    text = text.replace(r'\approx', '≈')
    text = text.replace(r'\neq', '≠')
    text = text.replace(r'\leq', '≤')
    text = text.replace(r'\geq', '≥')
    text = text.replace(r'\times', '×')
    text = text.replace(r'\rightarrow', '→')
    text = text.replace(r'\leftarrow', '←')
    
    # Remove simple commands like \textit{...}, \textbf{...}, \texttt{...}, \url{...}
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\url\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\href\{([^}]+)\}\{([^}]+)\}', r'\2 (\1)', text)
    text = re.sub(r'\\cite\{([^}]+)\}', r'[\1]', text)
    text = re.sub(r'\\ref\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\label\{([^}]+)\}', '', text)
    
    # Clean up quote marks
    text = text.replace('``', '"').replace("''", '"')
    text = text.replace('`', "'")
    
    # Replace LaTeX special accent notation if any
    # e.g. \'a -> á, \~n -> ñ
    accents = {
        r"\'a": "á", r"\'e": "é", r"\'i": "í", r"\'o": "ó", r"\'u": "ú",
        r"\'A": "Á", r"\'E": "É", r"\'I": "Í", r"\'O": "Ó", r"\'U": "Ú",
        r"\~n": "ñ", r"\~N": "Ñ", r'\"u': "ü", r'\"U': "Ü"
    }
    for k, v in accents.items():
        text = text.replace(k, v)
        
    return text

def parse_latex_table(table_text):
    rows = []
    # Find tabular body
    tabular_match = re.search(r'\\begin\{tabular\}\{([^\}]+)\}(.*?)\\end\{tabular\}', table_text, re.DOTALL)
    if not tabular_match:
        return None
    
    col_spec = tabular_match.group(1)
    body = tabular_match.group(2)
    
    # Split by rows
    lines = body.split(r'\\')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Skip rules
        if line.startswith(r'\toprule') or line.startswith(r'\midrule') or line.startswith(r'\bottomrule') or line.startswith(r'\hline'):
            continue
        # Split by columns
        cols = [clean_latex(c.strip()) for c in line.split('&')]
        # Filter out empty rows or lines that are just rules
        if all(c == '' for c in cols):
            continue
        rows.append(cols)
    return rows

def build_word():
    root_dir = r"C:\Users\renzo\amor-al-arte"
    art2_dir = os.path.join(root_dir, "articulo 2")
    latex_path = os.path.join(art2_dir, "fuente_latex", "articulo_scopus.tex")
    docx_path = os.path.join(art2_dir, "01_Informe_Articulo_Cientifico.docx")
    
    # Read LaTeX
    with open(latex_path, "r", encoding="utf-8") as f:
        latex_content = f.read()

    doc = Document()
    
    # Set Margins (2.54 cm = 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Choques Macroeconómicos y su Impacto en la Inflación y el Crecimiento Económico en Perú: Un Análisis SVAR para el Período 2000–2024\n")
    run.bold = True
    run.font.size = Pt(14)
    
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = author_p.add_run(
        "Estudiante de la Escuela Profesional de Ingeniería Económica\n"
        "Facultad de Ingeniería Económica, Universidad Nacional del Altiplano\n"
        "Puno, Perú\n\n"
        "Enlace al repositorio de Google Colab:\n"
        "https://colab.research.google.com/github/gaelrenzo/amor-al-arte/blob/main/articulo%202/04_Notebook_Google_Colab.ipynb\n"
    )
    run_auth.font.size = Pt(11)
    
    # Page break after Title Page
    doc.add_page_break()

    # Define some economics addition paragraphs to make it tie to Ingeniería Económica explicitly
    ing_econ_intro = (
        "En el contexto de la Ingeniería Económica y el análisis de políticas públicas, "
        "el estudio y modelamiento cuantitativo de las fluctuaciones macroeconómicas resulta vital. "
        "La Ingeniería Económica proporciona los marcos metodológicos para la evaluación de proyectos "
        "y decisiones de inversión tanto públicas como privadas, las cuales están inherentemente "
        "condicionadas por el entorno macroeconómico (tasa de interés, inflación y crecimiento). "
        "Por ende, estimar con precisión matemática los efectos y la persistencia de las perturbaciones "
        "estructurales mediante modelos de Vectores Autorregresivos Estructurales (SVAR) se convierte "
        "en una competencia analítica medular para el Ingeniero Economista. Este informe describe la "
        "dinámica de shocks de oferta, demanda, política fiscal y monetaria en la economía peruana, "
        "proporcionando herramientas empíricas y simulaciones rigurosas para la planificación económica."
    )
    
    ing_econ_method = (
        "Desde la perspectiva de la Ingeniería Económica aplicada a la macroeconomía, la estimación de "
        "multiplicadores y respuestas dinámicas es crucial para analizar la viabilidad financiera de "
        "las políticas gubernamentales. La identificación del modelo SVAR mediante restricciones de exclusión "
        "contemporánea permite aislar el efecto puro de cada choque estructural. El Ingeniero Economista utiliza "
        "estos resultados para evaluar la estabilidad de las tasas de interés y proyectar escenarios de "
        "costo-beneficio en inversiones públicas a mediano y largo plazo."
    )

    ing_econ_concl = (
        "Se concluye que el modelamiento econométrico SVAR estructurado representa una herramienta "
        "analítica fundamental para la Ingeniería Económica, integrando de manera sólida la teoría neokeynesiana "
        "de fluctuaciones con las técnicas estadísticas multivariantes de series de tiempo. Los hallazgos de este "
        "estudio respecto al impacto de la política fiscal y monetaria sirven como base cuantitativa para que los "
        "profesionales de la Ingeniería Económica diseñen políticas contracíclicas y evalúen el riesgo macroeconómico "
        "en la formulación y evaluación de proyectos nacionales."
    )

    # Parse sections from LaTeX
    lines = latex_content.split('\n')
    i = 0
    in_verbatim = False
    verbatim_lines = []
    current_section = ""

    while i < len(lines):
        line = lines[i]
        
        # Check verbatim blocks
        if r'\begin{verbatim}' in line:
            in_verbatim = True
            verbatim_lines = []
            i += 1
            continue
        elif r'\end{verbatim}' in line:
            in_verbatim = False
            # Add monospaced code block
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            code_text = "\n".join(verbatim_lines)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            # Add a light border/shading style if possible
            i += 1
            continue
            
        if in_verbatim:
            verbatim_lines.append(line)
            i += 1
            continue

        # Skip document class, packages, settings
        if any(line.startswith(x) for x in [r'\documentclass', r'\usepackage', r'\geometry', r'\hypersetup', r'\Declare', r'\begin{document}', r'\maketitle', r'\end{document}']):
            i += 1
            continue

        # Check section and subsection headers
        section_match = re.match(r'\\section\{([^\}]+)\}', line)
        section_star_match = re.match(r'\\section\*\{([^\}]+)\}', line)
        subsection_match = re.match(r'\\subsection\{([^\}]+)\}', line)
        subsection_star_match = re.match(r'\\subsection\*\{([^\}]+)\}', line)
        
        if section_match or section_star_match:
            title = clean_latex(section_match.group(1) if section_match else section_star_match.group(1))
            if current_section == "Conclusiones":
                # Add conclusion paragraph before transitioning
                p_concl = doc.add_paragraph()
                run_concl = p_concl.add_run(ing_econ_concl)
                run_concl.font.size = Pt(12)
                p_concl.paragraph_format.space_after = Pt(6)
                p_concl.paragraph_format.line_spacing = 1.5
                print("Appended concluding Ingeniería Económica paragraph.")
            current_section = title
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(title)
            run.bold = True
            run.font.size = Pt(13)
            
            # If we enter "Introducción", insert the Engineering Economics contextual paragraph!
            if "Introducción" in title:
                p = doc.add_paragraph()
                run_p = p.add_run(ing_econ_intro)
                run_p.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5

            i += 1
            continue
            
        elif subsection_match or subsection_star_match:
            title = clean_latex(subsection_match.group(1) if subsection_match else subsection_star_match.group(1))
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(title)
            run.bold = True
            run.font.size = Pt(12)

            # Insert Engineering Economics paragraph in Methodology
            if "Procedimiento econométrico" in title:
                p = doc.add_paragraph()
                run_p = p.add_run(ing_econ_method)
                run_p.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
                
            # If we are in "Resultados" subsections, let's insert the graphics!
            if "estabilidad del VAR" in title.lower():
                # Let's insert the stability plot
                img_path = os.path.join(art2_dir, "graficos", "graficos_estabilidad.png")
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run("Figura 1. Raíces inversas del polinomio característico del modelo VAR(4). Muestra la estabilidad dinámica del modelo estimado.")
                    run_cap.font.size = Pt(10)
                    run_cap.italic = True
                    
            elif "Impulso-Respuesta" in title:
                # Let's insert the IRF plots
                img_path1 = os.path.join(art2_dir, "graficos", "graficos_irf_all.png")
                img_path2 = os.path.join(art2_dir, "graficos", "graficos_irf_monetario.png")
                if os.path.exists(img_path1):
                    p_img1 = doc.add_paragraph()
                    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img1.add_run().add_picture(img_path1, width=Inches(5.0))
                    p_cap1 = doc.add_paragraph()
                    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap1 = p_cap1.add_run("Figura 2. Funciones Impulso-Respuesta acumuladas ante shocks estructurales de desviación estándar en el modelo SVAR.")
                    run_cap1.font.size = Pt(10)
                    run_cap1.italic = True
                    
                if os.path.exists(img_path2):
                    p_img2 = doc.add_paragraph()
                    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img2.add_run().add_picture(img_path2, width=Inches(5.0))
                    p_cap2 = doc.add_paragraph()
                    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap2 = p_cap2.add_run("Figura 3. Respuestas macroeconómicas específicas ante un choque de Tasa de Referencia de Política Monetaria.")
                    run_cap2.font.size = Pt(10)
                    run_cap2.italic = True
                    
            elif "Varianza del Error" in title:
                # Let's insert the FEVD plot
                img_path = os.path.join(art2_dir, "graficos", "graficos_fevd.png")
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run("Figura 4. Descomposición de la varianza del error de pronóstico de la inflación en un horizonte de 24 trimestres.")
                    run_cap.font.size = Pt(10)
                    run_cap.italic = True
                    
            elif "Histórica" in title:
                # Let's insert the historical decomposition plot
                img_path = os.path.join(art2_dir, "graficos", "graficos_descomposicion_historica.png")
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run("Figura 5. Descomposición histórica de shocks sobre la inflación peruana, destacando la contribución durante el período de pandemia (2020).")
                    run_cap.font.size = Pt(10)
                    run_cap.italic = True

            i += 1
            continue

        # Check tables
        if r'\begin{table}' in line:
            table_lines = []
            while i < len(lines) and r'\end{table}' not in lines[i]:
                table_lines.append(lines[i])
                i += 1
            table_lines.append(lines[i])
            table_block = "\n".join(table_lines)
            
            # Find caption
            caption_match = re.search(r'\\caption\{([^\}]+)\}', table_block)
            caption = clean_latex(caption_match.group(1)) if caption_match else "Tabla"
            
            # Add caption
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_cap.paragraph_format.space_before = Pt(6)
            p_cap.paragraph_format.space_after = Pt(2)
            p_cap.paragraph_format.keep_with_next = True
            run_cap = p_cap.add_run(caption)
            run_cap.bold = True
            run_cap.font.size = Pt(10)
            
            # Parse rows
            rows_data = parse_latex_table(table_block)
            if rows_data:
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.style = 'Table Grid'
                
                # Style and fill the table
                for r_idx, row in enumerate(rows_data):
                    for c_idx, val in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = val
                        
                        # Style text in cells
                        for p in cell.paragraphs:
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.0
                            for r in p.runs:
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(9.5)
                                if r_idx == 0:
                                    r.bold = True
                                    
                        # Set cell padding and subtle borders
                        set_cell_border(
                            cell,
                            top={"sz": 4, "val": "single", "color": "CCCCCC", "space": "0"},
                            bottom={"sz": 4, "val": "single", "color": "CCCCCC", "space": "0"},
                            left={"sz": 0, "val": "none"},
                            right={"sz": 0, "val": "none"}
                        )
                # Add space after table
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(6)
                p_space.paragraph_format.space_after = Pt(6)
                
            i += 1
            continue

        # Skip figure blocks (since we insert images manually in headings/subsections)
        if r'\begin{figure}' in line:
            while i < len(lines) and r'\end{figure}' not in lines[i]:
                i += 1
            i += 1
            continue

        # Skip bibliography environment markers
        if r'\begin{thebibliography}' in line or r'\end{thebibliography}' in line:
            i += 1
            continue
            
        # Parse bibliography items
        if line.strip().startswith(r'\bibitem'):
            # It's a reference list item
            ref_text = line.strip()
            # Extract content from \bibitem{label} text
            ref_match = re.search(r'\\bibitem\{[^}]+\}\s*(.*)', ref_text)
            if ref_match:
                content_ref = clean_latex(ref_match.group(1))
                # If reference is split across multiple lines, read them
                while i + 1 < len(lines) and not lines[i+1].strip().startswith(r'\bibitem') and not lines[i+1].strip().startswith(r'\end{') and not lines[i+1].strip().startswith(r'\section'):
                    i += 1
                    content_ref += " " + clean_latex(lines[i].strip())
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.27)
                p.paragraph_format.first_line_indent = Cm(-1.27)
                run = p.add_run(content_ref)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            i += 1
            continue



        # Add normal text paragraphs
        text = clean_latex(line.strip())
        if text:
            # Check if this is a subsection title or if it is part of body text
            # Some text might contain comments, let's ignore comments starting with %
            if not text.startswith('%'):
                p = doc.add_paragraph()
                p.add_run(text)
                
        i += 1

    # Add extra files in Annexes: Anexo B
    # Let's insert the diagnostic results in the Word document under Anexo B!
    for p in doc.paragraphs:
        if "Anexo B" in p.text:
            # We insert diagnostics content right after this paragraph
            # Read var_summary.txt and svar_summary.txt
            var_sum_path = os.path.join(art2_dir, "resultados", "var_summary.txt")
            svar_sum_path = os.path.join(art2_dir, "resultados", "svar_summary.txt")
            diag_path = os.path.join(art2_dir, "resultados", "diagnosticos_resumen.txt")
            
            content_to_insert = []
            if os.path.exists(diag_path):
                content_to_insert.append("Resumen de Pruebas de Diagnóstico del VAR(4):\n" + open(diag_path, "r", encoding="utf-8").read())
            if os.path.exists(svar_sum_path):
                content_to_insert.append("Resumen de Estimación de la Matriz B Estructural:\n" + open(svar_sum_path, "r", encoding="utf-8").read())
            if os.path.exists(var_sum_path):
                content_to_insert.append("Resumen Estadístico Detallado del VAR (Primeros Bloques):\n" + "\n".join(open(var_sum_path, "r", encoding="utf-8").read().splitlines()[:60]))
                
            for block in content_to_insert:
                p_block = doc.add_paragraph()
                p_block.paragraph_format.line_spacing = 1.0
                p_block.paragraph_format.space_before = Pt(6)
                p_block.paragraph_format.space_after = Pt(2)
                p_block.paragraph_format.left_indent = Cm(0.5)
                run_block = p_block.add_run(block)
                run_block.font.name = 'Courier New'
                run_block.font.size = Pt(8.0)
            break

    # Save document
    doc.save(docx_path)
    print("Successfully built the expanded DOCX report!")

if __name__ == "__main__":
    build_word()
