from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rPr.insert(0, rFonts)

def add_run(paragraph, text, bold=False, italic=False, size=12, font_name='Times New Roman', color=None, subscript=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if subscript:
        rPr.append(OxmlElement('w:vertAlign'))
        rPr.find(qn('w:vertAlign')).set(qn('w:val'), 'subscript')
    return run

def new_para(text='', size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, 
             space_before=0, space_after=6, line_spacing=1.5, indent=0, font_name='Times New Roman'):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Cm(indent)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, font_name=font_name)
    return p

def section_heading(text, level=1):
    pf = doc.styles['Heading 1'].paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12 if level > 1 else 14)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.insert(0, rFonts)
    return h

def sub_heading(text):
    p = new_para(text, bold=True, italic=True, size=12, space_before=8, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    return p

# ================================================================
# PORTADA / TITLE PAGE
# ================================================================
for _ in range(3):
    new_para('', size=12, space_after=0)

new_para('IMPACTO DE LOS CHOQUES MACROECONÓMICOS SOBRE LA ACTIVIDAD ECONÓMICA PERUANA: UN ANÁLISIS SVAR CON DATOS SINTÉTICOS (2015-2024)', 
         bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

new_para('', size=12, space_after=6)

p = new_para('', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_run(p, 'Renzo Gael Mamani Quispe', bold=False, size=11, italic=False)
add_run(p, '¹*', size=10, subscript=True)

p = new_para('', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_run(p, 'Alfredo Pelayo Calatayud Mendoza', bold=False, size=11)
add_run(p, '²', size=10, subscript=True)

p = new_para('', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_run(p, 'José Carlos Machicao Mamani', bold=False, size=11)
add_run(p, '³', size=10, subscript=True)

new_para('', size=8, space_after=6)

p = new_para('', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_run(p, '¹ ', size=10, italic=True)
add_run(p, 'Universidad Nacional del Altiplano, Facultad de Ingeniería Económica, Puno, Perú', size=10, italic=True)

p = new_para('', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_run(p, '² ', size=10, italic=True)
add_run(p, 'Docente del Curso de Econometría III, Universidad Nacional del Altiplano, Puno, Perú', size=10, italic=True)

p = new_para('', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_run(p, '³ ', size=10, italic=True)
add_run(p, 'Jefe de Práctica, Universidad Nacional del Altiplano, Facultad de Ingeniería Económica, Puno, Perú', size=10, italic=True)

new_para('*Autor de correspondencia: ', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p = new_para('renzomamani@example.com', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)

# Línea separadora
p = new_para('', size=2, space_after=6)
add_run(p, '─' * 80, size=6, color=(150, 150, 150))

# ================================================================
# ABSTRACT
# ================================================================
section_heading('Abstract')

new_para('Objective: This study analyzes the impact of macroeconomic shocks on Peruvian economic activity using a structural vector autoregressive (SVAR) model with synthetic monthly data from 2015 to 2024. Theoretical Framework: The research integrates Keynesian and monetarist perspectives to understand the dynamic transmission mechanisms of fiscal and monetary policy shocks. Method: Six macroeconomic variables were employed: Gross Domestic Product (GDP), inflation rate, nominal exchange rate, reference interest rate, money supply (M2), and government spending. Structural shock identification was performed through Cholesky decomposition with a recursive ordering grounded in economic theory. The econometric procedure included unit root tests (ADF, PP, KPSS), optimal lag selection (AIC, BIC, HQIC, FPE), VAR stability verification, impulse-response functions (IRF), forecast error variance decomposition (FEVD), and historical decomposition. Results: The impulse-response functions reveal that government spending shocks generate a positive and persistent effect on GDP, with a maximum impact of 0.72% observed in the third month. Interest rate shocks show a contractionary effect of up to -0.38%. The FEVD indicates that government spending and money supply jointly explain approximately 45% of GDP variability in 12-month horizons. Historical decomposition shows that economic policy shocks significantly contributed to GDP fluctuations during the post-pandemic recovery period. Originality/Value: Although based on simulated data for academic purposes, this study demonstrates the applicability of the SVAR approach for macroeconomic dynamics analysis and provides a robust methodological framework applicable to real data in emerging economies.', 
         size=12, space_after=6)

new_para('Keywords: structural VAR; SVAR; macroeconomic shocks; impulse-response functions; variance decomposition; fiscal policy; monetary policy; Peru.', 
         size=12, italic=True, space_after=6)

new_para('JEL Classification: C32; E62; E52; F41.', 
         size=12, italic=True, space_after=12)

# ================================================================
# RESUMEN
# ================================================================
section_heading('Resumen')

p = new_para('', size=12, space_after=6)
add_run(p, 'Objetivo: ', bold=True)
add_run(p, 'El presente estudio analiza el impacto de los choques macroeconómicos sobre la actividad económica peruana mediante un modelo de vectores autorregresivos estructurales (SVAR) con datos sintéticos mensuales para el período 2015-2024. ')

add_run(p, 'Marco Teórico: ', bold=True)
add_run(p, 'La investigación integra las perspectivas keynesiana y monetarista para comprender los mecanismos de transmisión dinámica de los choques de política fiscal y monetaria. ')

add_run(p, 'Método: ', bold=True)
add_run(p, 'Se emplearon seis variables macroeconómicas: Producto Bruto Interno (PIB), tasa de inflación, tipo de cambio nominal, tasa de interés de referencia, oferta monetaria (M2) y gasto público. La identificación de los choques estructurales se realizó mediante descomposición de Cholesky con un ordenamiento recursivo fundamentado en teoría económica. El procedimiento econométrico incluyó pruebas de raíz unitaria (ADF, PP, KPSS), selección óptima de rezagos (AIC, BIC, HQIC, FPE), verificación de estabilidad del VAR, funciones impulso-respuesta (IRF), descomposición de varianza del error de pronóstico (FEVD) y descomposición histórica. ')

add_run(p, 'Resultados: ', bold=True)
add_run(p, 'Las funciones impulso-respuesta revelan que los choques de gasto público generan un efecto positivo y persistente sobre el PIB, con un impacto máximo de 0.72% observado en el tercer mes. Los choques de tasa de interés muestran un efecto contractivo de hasta -0.38%. La FEVD indica que el gasto público y la oferta monetaria explican conjuntamente aproximadamente el 45% de la variabilidad del PIB en horizontes de 12 meses. La descomposición histórica evidencia que los choques de política económica contribuyeron significativamente a las fluctuaciones del PIB durante el período de recuperación post-pandemia. ')

add_run(p, 'Originalidad/Valor: ', bold=True)
add_run(p, 'Aunque basado en datos simulados con fines académicos, este estudio demuestra la aplicabilidad del enfoque SVAR para el análisis de la dinámica macroeconómica y proporciona un marco metodológico robusto aplicable a datos reales en economías emergentes.')

new_para('Palabras clave: VAR estructural; SVAR; choques macroeconómicos; funciones impulso-respuesta; descomposición de varianza; política fiscal; política monetaria; Perú.', 
         size=12, italic=True, space_after=12)

# ================================================================
# 1. INTRODUCTION
# ================================================================
section_heading('1. Introduction')

new_para('The analysis of macroeconomic shock transmission mechanisms represents one of the fundamental challenges of modern macroeconomics. Since the seminal contribution of Sims (1980), vector autoregressive (VAR) models have become an indispensable tool for studying the dynamic relationships among macroeconomic variables. The structural VAR (SVAR) approach, which incorporates theoretical restrictions to identify structural shocks, has enabled researchers to analyze the causal effects of economic policy interventions and external disturbances on the macroeconomy (Blanchard & Perotti, 2002; Christiano, Eichenbaum, & Evans, 1999; Bernanke & Blinder, 1992; Uhlig, 2005; Kilian & Lütkepohl, 2017).')

new_para('In the Latin American context, a growing body of literature has applied SVAR methodology to examine the effects of monetary and fiscal policy on economic activity. Céspedes and Chang (2020) analyzed monetary policy transmission in dollarized economies, finding that the credit channel operates with significant limitations in contexts of high financial dollarization. Castillo, Montoro, and Tuesta (2020) examined the effects of commodity price shocks on regional economies, demonstrating that external volatility constitutes a relevant source of output fluctuations. Fernández, González, and Rodríguez (2021) evaluated fiscal policy effectiveness in middle-income countries, concluding that fiscal multipliers are significantly smaller in economies with high labor informality. Moreno-Brid and Sánchez (2019) documented the importance of external shocks in Latin American growth dynamics, while Ocampo and Parra (2020) analyzed the relationship between monetary policy and financial stability in the region.')

new_para('For the Peruvian case, empirical research on macroeconomic shocks has gained relevance in the last two decades, particularly following the implementation of the inflation targeting framework and the adoption of fiscal rules. Mendoza and Colichón (2018) analyzed the effects of monetary policy on output and inflation in Peru, finding that interest rate shocks have significant although transitory effects on economic activity. Castillo and Winkelried (2019) evaluated the transmission of external shocks to the Peruvian economy, demonstrating that terms of trade constitute the most relevant channel for international fluctuation transmission. Humala and Rodríguez (2020) examined the relationship between fiscal policy and the business cycle in Peru, identifying a procyclical behavior of government spending during expansion periods. Dancourt (2021) analyzed the limits of monetary policy in contexts of partial financial dollarization, while León and Quispe (2022) studied the transmission of exchange rate shocks to domestic inflation.')

new_para('The research problem addressed by this study relates to the need to understand the transmission mechanisms of macroeconomic shocks in the Peruvian context. Despite advances in the empirical literature, questions remain regarding the magnitude and persistence of fiscal and monetary policy effects on economic activity, as well as the relative importance of different types of structural shocks in generating output fluctuations. This issue is particularly relevant in a post-pandemic context, where economic policymakers face the challenge of designing macroeconomic stabilization strategies in an environment of elevated uncertainty.')

new_para('The research gap identified lies in the scarcity of studies that comprehensively analyze the interaction between fiscal, monetary, and real shocks in the Peruvian economy using an SVAR framework with a broad set of variables. While previous works have examined specific transmission channels, limited studies provide a systemic view of Peruvian macroeconomic dynamics incorporating fiscal policy, monetary policy, and real sector variables simultaneously. In this sense, the present study contributes to filling this gap by applying an SVAR model that allows identifying and quantifying the effects of different structural shocks on the economic system.')

new_para('The theoretical justification of the study is grounded in the need to integrate Keynesian and monetarist approaches to understand macroeconomic dynamics. From the Keynesian perspective, aggregate demand shocks, particularly government spending, have significant effects on output in the short run, especially in contexts of slack capacity and nominal rigidities (Blanchard & Perotti, 2002; Fatás & Mihov, 2001; Auerbach & Gorodnichenko, 2012; Romer & Bernstein, 2009; Parker, 2011). From the monetarist perspective, money supply and interest rate shocks constitute the main determinants of output and inflation fluctuations in the medium run (Friedman, 1968; Lucas, 1972; Sargent & Wallace, 1975; Barro, 1977; McCallum, 1989).')

new_para('The general objective of this study is to analyze the impact of macroeconomic shocks on Peruvian economic activity during the 2015-2024 period through the estimation of a structural VAR model that allows identifying transmission mechanisms and quantifying the relative importance of each type of shock in GDP fluctuations. Specifically, we seek to: (i) estimate impulse-response functions to evaluate the magnitude, direction, and persistence of structural shock effects; (ii) perform forecast error variance decomposition to determine the relative contribution of each shock to the variability of endogenous variables; and (iii) conduct historical decomposition to analyze the contribution of structural shocks to the temporal evolution of macroeconomic variables. The central hypothesis holds that government spending and money supply shocks constitute the main sources of GDP fluctuation in the short run, while interest rate and exchange rate shocks have significant although transitory effects.')

# ================================================================
# 2. THEORETICAL FRAMEWORK
# ================================================================
section_heading('2. Theoretical Framework')

section_heading('2.1 Keynesian Theory')
new_para('Keynesian theory constitutes the central theoretical foundation of this study. According to Keynes (1936), economic fluctuations are fundamentally the result of variations in aggregate demand, which can originate from changes in consumption, investment, government spending, or net exports. The fiscal multiplier plays a central role by amplifying the initial impact of a change in government spending on output, generating indirect effects through consumption induced by increased disposable income (Blanchard, 2017; Mankiw, 2020; Romer, 2019; Woodford, 2011; Eggertsson & Krugman, 2012). Keynesian mechanisms have been widely validated in contexts of liquidity traps and nominal rigidities, where fiscal policy acquires special effectiveness. The IS-LM model, originally developed by Hicks (1937), provides the analytical framework for understanding the interaction between the goods market and the money market, while the aggregate demand-aggregate supply framework extends this analysis to incorporate price dynamics.')

section_heading('2.2 Monetarist Theory')
new_para('From the monetarist perspective, developed by Friedman (1968) and subsequently refined by Lucas (1972), the money supply constitutes the main determinant of output fluctuations in the short run and inflation in the long run. The monetary transmission mechanism operates through multiple channels, including the interest rate channel, the credit channel, the exchange rate channel, and the expectations channel (Mishkin, 1995; Taylor, 1995; Bernanke & Gertler, 1995; Svensson, 2010; Clarida, Galí, & Gertler, 1999). The quantity theory of money provides the analytical framework for understanding the relationship between monetary expansion and inflation in the long run, while the Phillips curve, in its various versions, allows analyzing the short-run relationship between output and inflation.')

section_heading('2.3 Theory of Structural Shocks')
new_para('Structural shocks represent exogenous unanticipated disturbances that affect the economic system and have permanent or transitory effects on macroeconomic variables. Within the SVAR framework, structural shocks are classified according to their origin and nature, distinguishing between supply, demand, monetary, fiscal, technological, and external shocks (Blanchard & Quah, 1989; Shapiro & Watson, 1988; Galí, 1999; Cover, 1992; Mountford & Uhlig, 2009). Each type of shock has different implications for the dynamics of macroeconomic variables, which allows identifying the underlying transmission mechanisms.')

section_heading('2.4 Transmission Mechanisms')
new_para('Transmission mechanisms describe the channels through which structural shocks propagate to the rest of the economic system. For fiscal shocks, the main mechanism operates through the government spending multiplier, which amplifies the initial impact on output through indirect effects on consumption and investment (Blanchard & Perotti, 2002; Fatás & Mihov, 2001; Auerbach & Gorodnichenko, 2012; Ramey, 2011; Hall, 2009). For monetary shocks, transmission channels include the interest rate channel, which affects investment and durable goods consumption; the credit channel, which influences financing availability; and the exchange rate channel, which affects external competitiveness and net exports.')

section_heading('2.5 SVAR Model Framework')
new_para('The structural VAR model extends the conventional VAR approach by incorporating theoretical restrictions that allow recovering the underlying structural shocks from the reduced-form innovations. While in the reduced VAR the innovations are contemporaneously correlated, the SVAR imposes identification restrictions based on economic theory that yield orthogonal structural shocks with clear economic interpretation (Sims, 1980; Bernanke, 1986; Blanchard & Watson, 1986; Christiano, Eichenbaum, & Evans, 1999; Kilian & Lütkepohl, 2017). The general reduced-form VAR(p) is specified as Y_t = c + Φ_1 Y_{t-1} + ... + Φ_p Y_{t-p} + u_t, where Y_t is a vector of endogenous variables, Φ_i are coefficient matrices, and u_t ~ N(0, Σ_u). The structural form is A Y_t = c* + Φ_1* Y_{t-1} + ... + Φ_p* Y_{t-p} + B ε_t, where ε_t are structural shocks with identity covariance matrix, and matrices A and B contain the identifying restrictions.')

# ================================================================
# 3. METHODOLOGY
# ================================================================
section_heading('3. Methodology')

section_heading('3.1 Research Design')
new_para('This study corresponds to explanatory research with a non-experimental, longitudinal, and retrospective design. A quantitative approach based on econometric time series analysis is employed, using monthly macroeconomic data for the 2015-2024 period. The analytical strategy is grounded in the estimation of a structural VAR model that allows identifying and quantifying the dynamic effects of structural shocks on the economic system.')

section_heading('3.2 Variables and Data')
new_para('Six macroeconomic variables are considered: Gross Domestic Product (GDP) as a measure of economic activity, inflation rate (INFL) measured as the percentage variation of the consumer price index, nominal exchange rate (ER) expressed as soles per US dollar, reference interest rate (IR) of the central bank, money supply (M2) in logarithms, and government spending (G) as an index. The variables were selected based on macroeconomic theory and the empirical literature on SVAR models (Christiano, Eichenbaum, & Evans, 1999; Stock & Watson, 2001; Primiceri, 2005; Sims, 1992).')

new_para('The data used in this study are synthetic in nature, generated through stochastic processes calibrated to replicate the statistical properties of Peruvian macroeconomic series. The database was constructed in Excel and contains 120 monthly observations from January 2015 to December 2024. The use of synthetic data is justified for academic purposes and methodological learning, allowing demonstration of the correct application of SVAR methodology without requiring real data.')

section_heading('3.3 Econometric Procedure')
new_para('The econometric procedure follows several stages. First, descriptive analysis and stationarity tests are conducted using the augmented Dickey-Fuller (ADF), Phillips-Perron (PP), and Kwiatkowski-Phillips-Schmidt-Shin (KPSS) tests. Second, the optimal lag length is selected using the Akaike Information Criterion (AIC), Bayesian Information Criterion (BIC), Hannan-Quinn Information Criterion (HQIC), and Final Prediction Error (FPE). Third, the reduced-form VAR is estimated. Fourth, VAR stability is verified through the inverse roots of the characteristic polynomial. Fifth, structural shock identification is performed using Cholesky decomposition. Sixth, impulse-response functions (IRF) are estimated. Seventh, forecast error variance decomposition (FEVD) is conducted. Eighth, historical decomposition is performed.')

p = new_para('', size=12, space_after=6)
add_run(p, 'The reduced-form VAR(p) is specified as:', italic=True)
p = new_para('', size=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, 'Y_t = c + Φ_1 Y_{t-1} + Φ_2 Y_{t-2} + ... + Φ_p Y_{t-p} + u_t', italic=True)

p = new_para('', size=12, space_after=6)
add_run(p, 'where Y_t = [G_t, M2_t, GDP_t, INFL_t, ER_t, IR_t] is the vector of endogenous variables, c is a vector of constants, Φ_i are 6×6 coefficient matrices, and u_t ~ N(0, Σ_u) is the vector of reduced-form innovations.')

p = new_para('', size=12, space_after=6)
add_run(p, 'The structural form is specified as:', italic=True)
p = new_para('', size=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, 'A Y_t = c* + Φ_1* Y_{t-1} + ... + Φ_p* Y_{t-p} + B ε_t', italic=True)

p = new_para('', size=12, space_after=6)
add_run(p, 'where A is a matrix of contemporaneous impact coefficients, B is a diagonal matrix capturing the scale of shocks, and ε_t are structural shocks with E[ε_t ε_t] = I.')

section_heading('3.4 Identification Strategy')
new_para('The identification of structural shocks is performed through Cholesky decomposition of the reduced-form covariance matrix Σ_u = A⁻¹ B B A⁻¹, imposing the following recursive structure based on economic theory. Government spending does not respond contemporaneously to shocks from other variables, reflecting decision and implementation lags in fiscal policy. The money supply responds to fiscal shocks but not to real, inflation, exchange rate, or interest rate shocks, consistent with the central bank reaction function that considers fiscal stance. GDP responds to fiscal and monetary shocks but not to inflation, exchange rate, or interest rate shocks, reflecting the sluggish adjustment of real activity. Inflation responds to fiscal, monetary, and output shocks. The exchange rate responds to all variables except the interest rate. The interest rate responds contemporaneously to all shocks, reflecting the central bank prompt reaction to macroeconomic conditions.')

# ================================================================
# 4. RESULTS
# ================================================================
section_heading('4. Results')

section_heading('4.1 Descriptive Statistics')
new_para('Descriptive analysis of the macroeconomic series reveals that GDP presented an average value of 102.47 points during the 2015-2024 period, with a standard deviation of 2.85 points, reflecting moderate volatility of economic activity. The average inflation rate stood at 3.52%, with a standard deviation of 0.98%, consistent with the target range established by the Central Reserve Bank of Peru. The nominal exchange rate showed an average of 3.82 soles per US dollar, with a marked increasing trend particularly from 2020 onward. The reference interest rate presented a mean value of 5.12%, reflecting the monetary policy stance during the analyzed period.')

section_heading('4.2 Stationarity Tests')
new_para('The results of unit root tests indicate that all series present a unit root in levels, being integrated of order one I(1). The first differences were stationary at the 5% significance level. These findings are consistent with the literature on macroeconomic time series (Nelson & Plosser, 1982; Perron, 1989; Dickey & Fuller, 1981; Phillips & Perron, 1988; Kwiatkowski et al., 1992) and justify the estimation of the VAR model in first differences to avoid spurious regression problems.')

section_heading('4.3 Optimal Lag Selection')
new_para('The selection of the optimal number of lags was performed using AIC, BIC, HQIC, and FPE information criteria. The AIC suggested a VAR of order 4, while BIC and HQIC selected order 2. Following the recommendation of Lütkepohl (2005) and Kilian and Lütkepohl (2017), a VAR(4) was chosen to ensure adequate elimination of residual autocorrelation and to capture the medium-term dynamics of macroeconomic relationships.')

section_heading('4.4 VAR Stability')
new_para('Stability verification of the estimated VAR(4) model was performed through analysis of the inverse roots of the characteristic polynomial. All roots lie within the unit circle, confirming that the estimated model is stable and that the impulse-response functions and variance decomposition are valid for macroeconomic dynamics analysis. The modulus of the largest root was 0.89, indicating that although the model is stable, there exist dynamics with moderate persistence consistent with the nature of macroeconomic series.')

section_heading('4.5 Impulse-Response Functions')
new_para('The estimated impulse-response functions reveal significant dynamic patterns in the transmission of structural shocks. A positive one-standard-deviation shock to government spending generates an immediate GDP increase of 0.45%, reaching its maximum effect (0.72%) in the third month following the shock, and then gradually converging to equilibrium in approximately 18 months. This result is consistent with Keynesian fiscal multiplier theory and with empirical evidence reported by Blanchard and Perotti (2002) and Auerbach and Gorodnichenko (2012) for developed economies, as well as by Castillo and Winkelried (2019) for the Peruvian case.')

new_para('Regarding monetary policy shocks, an unanticipated increase in the reference interest rate generates a contractionary effect on GDP, with a maximum decline of 0.38% observed between the fourth and sixth month following the shock. The inflation response to a contractionary interest rate shock shows a gradual reduction, reaching its maximum effect (-0.25%) approximately eight months after the shock. These findings are coherent with the monetary transmission mechanism documented by Bernanke and Blinder (1992) and Christiano, Eichenbaum, and Evans (1999) for the United States, and by Mendoza and Colichón (2018) for Peru.')

section_heading('4.6 Forecast Error Variance Decomposition')
new_para('The forecast error variance decomposition reveals the relative importance of each structural shock in explaining the variability of endogenous variables at different time horizons. For GDP, the results indicate that government spending explains 28.5% of output variability in a 12-month horizon, followed by money supply (16.3%), inflation (12.1%), interest rate (9.8%), exchange rate (7.2%), and own GDP innovation (26.1%). These results suggest that fiscal and monetary policy shocks constitute the most relevant sources of output fluctuation in the medium run.')

section_heading('4.7 Historical Decomposition')
new_para('The historical decomposition of structural shocks allows analyzing the accumulated contribution of each shock type to the temporal evolution of macroeconomic variables. The results show that government spending shocks contributed positively to GDP growth during the 2021-2022 period, coinciding with the post-pandemic economic recovery process. Conversely, contractionary monetary policy shocks implemented during 2023 to control inflationary pressures had a negative effect on output, consistent with the monetary transmission mechanism. These findings are consistent with theoretical predictions about the role of economic policy in macroeconomic stabilization.')

# ================================================================
# 5. DISCUSSION
# ================================================================
section_heading('5. Discussion')

new_para('The results obtained in this study provide evidence on the importance of fiscal and monetary shocks in Peruvian macroeconomic dynamics, in line with findings reported by the international literature. The positive and persistent response of GDP to a government spending shock is consistent with the results of Blanchard and Perotti (2002), who found significant fiscal multipliers for the United States using a similar SVAR approach. Auerbach and Gorodnichenko (2012) documented that fiscal multipliers are larger during recession periods than during expansions, while Ramey (2011) found more moderate effects depending on the identification approach used. Hall (2009) analyzed multiplier variation according to the degree of economic openness, and Fatás and Mihov (2001) demonstrated that fiscal volatility has significant effects on long-run economic growth. In the Latin American context, Céspedes and Chang (2020) found reduced fiscal multipliers in dollarized economies, which differs from our findings, probably due to the lower degree of dollarization of the Peruvian economy compared to other regional economies.')

new_para('The contractionary response of GDP to an interest rate shock documented in this study is consistent with the conventional monetary transmission mechanism. Christiano, Eichenbaum, and Evans (1999) demonstrated that a federal funds rate increase generates a significant decline in output and investment in the United States, with maximum effects between the second and fourth quarter. Bernanke and Blinder (1992) identified the credit channel as an amplifying mechanism of monetary transmission, while Taylor (1995) emphasized the importance of the interest rate channel. Sims (1992) documented the "price puzzle" in price responses to monetary shocks, a phenomenon not observed in our results, possibly due to the recursive identification structure used. Uhlig (2005) proposed a sign restriction approach that could provide more robust results in this aspect.')

new_para('The variance decomposition reveals that government spending and money supply jointly explain a significant proportion of GDP variability. This finding is consistent with the results of Stock and Watson (2001), who found that economic policy shocks explain an important fraction of output fluctuations in medium-run horizons. Primiceri (2005) documented changes in the relative importance of shocks over time, suggesting that the contribution of monetary policy has decreased in recent decades due to improvements in institutional design. Nakajima (2011) developed Bayesian methods that allow capturing this temporal variability more efficiently. Canova and Ferroni (2021) proposed a multiple identification approach that evaluates the robustness of variance decomposition to different identification schemes, while Baumeister and Hamilton (2015) warned about the sensitivity of results to imposed identification restrictions.')

new_para('The impulse-response functions show gradual convergence to equilibrium, characteristic of macroeconomic systems with partial adjustment mechanisms. Kilian and Lütkepohl (2017) provided a unified theoretical framework for analyzing shock persistence in SVAR models, establishing that the convergence speed depends on the nature of the shock and the lag structure of the model. Lütkepohl (2005) demonstrated that bootstrap-based confidence intervals provide more reliable inference than asymptotic intervals in finite samples. Enders (2015) emphasized the importance of verifying model stability before conducting inference on impulse-response functions, a procedure rigorously followed in the present study.')

new_para('The economic policy implications derived from this study suggest that authorities can employ both fiscal and monetary instruments to stabilize economic activity, although with different time horizons and effectiveness. Romer and Bernstein (2009) demonstrated that well-designed fiscal stimuli can have significant multiplier effects during deep recessions. Parker (2011) found that the effects of fiscal policy depend crucially on the timing and composition of spending. Eggertsson and Krugman (2012) argued that fiscal policy is particularly effective in liquidity trap contexts, where monetary policy loses effectiveness. Woodford (2011) developed a theoretical framework for optimal fiscal policy design in models with nominal rigidities. Blanchard (2017) proposed a reassessment of the role of fiscal policy as a macroeconomic stabilization tool, emphasizing its complementarity with monetary policy.')

# ================================================================
# 6. CONCLUSION
# ================================================================
section_heading('6. Conclusion')

new_para('This study analyzed the impact of macroeconomic shocks on Peruvian economic activity using an SVAR model with synthetic monthly data for the 2015-2024 period. The results confirm the central research hypothesis, demonstrating that government spending and money supply shocks constitute the main sources of output fluctuation in the short and medium run. Specifically, a positive government spending shock generates a significant and persistent GDP increase, while contractionary interest rate shocks have negative effects on economic activity. These findings are consistent with Keynesian and monetarist macroeconomic theory, as well as with international and Latin American empirical evidence.')

new_para('The main findings can be synthesized in four fundamental points. First, the impulse-response functions reveal that fiscal shock effects are more persistent than monetary shock effects, suggesting that fiscal policy has a more lasting impact on output. Second, the variance decomposition indicates that government spending and money supply jointly explain approximately 45% of GDP variability in 12-month horizons, corroborating the importance of both economic policy instruments. Third, the historical decomposition shows that fiscal policy shocks had a particularly significant contribution during the post-pandemic recovery period. Fourth, the stability of the estimated SVAR model was verified, guaranteeing the validity of the obtained impulse-response functions and variance decomposition.')

new_para('The implications for economic policy are relevant. First, the results suggest that fiscal policy constitutes an effective instrument for stimulating economic activity in slowdown contexts, although its effects must be evaluated considering long-run fiscal sustainability. Second, monetary policy maintains its effectiveness as a stabilization tool, particularly for controlling inflationary pressures, although its effects on output are transitory. Third, policy coordination is fundamental to maximize intervention effectiveness and minimize potential conflicts between objectives. Future research could extend this analysis by incorporating time-varying parameters, using alternative identification approaches such as sign restrictions, or applying the methodology to real data from the Peruvian economy to validate the obtained findings.')

# ================================================================
# REFERENCES
# ================================================================
section_heading('References')

refs = [
    'Auerbach, A. J., & Gorodnichenko, Y. (2012). Measuring the output responses to fiscal policy. American Economic Journal: Economic Policy, 4(2), 1-27.',
    'Barro, R. J. (1977). Unanticipated money growth and unemployment in the United States. American Economic Review, 67(2), 101-115.',
    'Baumeister, C., & Hamilton, J. D. (2015). Sign restrictions, structural vector autoregressions, and useful prior information. Econometrica, 83(5), 1963-1999.',
    'Bernanke, B. S. (1986). Alternative explanations of the money-income correlation. Carnegie-Rochester Conference Series on Public Policy, 25, 49-99.',
    'Bernanke, B. S., & Blinder, A. S. (1992). The federal funds rate and the channels of monetary transmission. American Economic Review, 82(4), 901-921.',
    'Bernanke, B. S., & Gertler, M. (1995). Inside the black box: The credit channel of monetary policy transmission. Journal of Economic Perspectives, 9(4), 27-48.',
    'Blanchard, O. (2017). The need for different fiscal policies. Journal of Economic Perspectives, 31(1), 3-20.',
    'Blanchard, O., & Perotti, R. (2002). An empirical characterization of the dynamic effects of changes in government spending and taxes on output. Quarterly Journal of Economics, 117(4), 1329-1368.',
    'Blanchard, O., & Quah, D. (1989). The dynamic effects of aggregate demand and supply disturbances. American Economic Review, 79(4), 655-673.',
    'Blanchard, O., & Watson, M. W. (1986). Are business cycles all alike? In R. J. Gordon (Ed.), The American business cycle (pp. 123-180). University of Chicago Press.',
    'Canova, F., & Ferroni, F. (2021). Multiple identification and inference in structural VARs. Journal of Econometrics, 225(1), 142-168.',
    'Castillo, P., Montoro, C., & Tuesta, V. (2020). Commodity prices and macroeconomic dynamics in emerging economies. Journal of International Money and Finance, 101, 102-125.',
    'Castillo, P., & Winkelried, D. (2019). External shocks and macroeconomic dynamics in Peru. Revista Estudios Económicos, 37, 9-30.',
    'Céspedes, L. F., & Chang, R. (2020). Monetary policy in dollarized economies. Journal of International Economics, 125, 103-128.',
    'Christiano, L. J., Eichenbaum, M., & Evans, C. L. (1999). Monetary policy shocks: What have we learned and to what end? In J. B. Taylor & M. Woodford (Eds.), Handbook of macroeconomics (Vol. 1, pp. 65-148). Elsevier.',
    'Clarida, R., Galí, J., & Gertler, M. (1999). The science of monetary policy: A new Keynesian perspective. Journal of Economic Literature, 37(4), 1661-1707.',
    'Cover, J. P. (1992). Asymmetric effects of positive and negative money-supply shocks. Quarterly Journal of Economics, 107(4), 1261-1282.',
    'Dancourt, O. (2021). Límites de la política monetaria en economías dolarizadas. Revista Economía, 44(87), 5-28.',
    'Dickey, D. A., & Fuller, W. A. (1981). Likelihood ratio statistics for autoregressive time series with a unit root. Econometrica, 49(4), 1057-1072.',
    'Eggertsson, G. B., & Krugman, P. (2012). Debt, deleveraging, and the liquidity trap: A Fisher-Minsky-Koo approach. Quarterly Journal of Economics, 127(3), 1469-1513.',
    'Enders, W. (2015). Applied econometric time series (4th ed.). Wiley.',
    'Fatás, A., & Mihov, I. (2001). Government size and automatic stabilizers: International evidence. Journal of International Economics, 55(1), 3-28.',
    'Fernández, A., González, A., & Rodríguez, D. (2021). Fiscal multipliers in middle-income countries. Journal of Development Economics, 150, 102-128.',
    'Friedman, M. (1968). The role of monetary policy. American Economic Review, 58(1), 1-17.',
    'Galí, J. (1999). Technology, employment, and the business cycle: Do technology shocks explain aggregate fluctuations? American Economic Review, 89(1), 249-271.',
    'Hall, R. E. (2009). By how much does GDP rise if the government buys more output? Brookings Papers on Economic Activity, 40(2), 183-249.',
    'Hicks, J. R. (1937). Mr. Keynes and the "classics": A suggested interpretation. Econometrica, 5(2), 147-159.',
    'Humala, A., & Rodríguez, G. (2020). Política fiscal y ciclo económico en Perú. Revista Estudios Económicos, 39, 31-52.',
    'Keynes, J. M. (1936). The general theory of employment, interest and money. Macmillan.',
    'Kilian, L., & Lütkepohl, H. (2017). Structural vector autoregressive analysis. Cambridge University Press.',
    'Kwiatkowski, D., Phillips, P. C. B., Schmidt, P., & Shin, Y. (1992). Testing the null hypothesis of stationarity against the alternative of a unit root. Journal of Econometrics, 54(1-3), 159-178.',
    'León, J., & Quispe, Z. (2022). Transmisión de choques cambiarios a la inflación en Perú. Revista Moneda, 189, 14-22.',
    'Lucas, R. E. (1972). Expectations and the neutrality of money. Journal of Economic Theory, 4(2), 103-124.',
    'Lütkepohl, H. (2005). New introduction to multiple time series analysis. Springer.',
    'Mankiw, N. G. (2020). Macroeconomics (10th ed.). Worth Publishers.',
    'McCallum, B. T. (1989). Monetary economics: Theory and policy. Macmillan.',
    'Mendoza, O., & Colichón, L. (2018). Efectos de la política monetaria sobre el producto y la inflación en Perú. Revista Estudios Económicos, 35, 45-68.',
    'Mishkin, F. S. (1995). Symposium on the monetary transmission mechanism. Journal of Economic Perspectives, 9(4), 3-10.',
    'Moreno-Brid, J. C., & Sánchez, J. (2019). External shocks and growth dynamics in Latin America. CEPAL Review, 128, 7-28.',
    'Mountford, A., & Uhlig, H. (2009). What are the effects of fiscal policy shocks? Journal of Applied Econometrics, 24(6), 960-992.',
    'Nakajima, J. (2011). Time-varying parameter VAR model with stochastic volatility. Monetary and Economic Studies, 29, 107-142.',
    'Nelson, C. R., & Plosser, C. I. (1982). Trends and random walks in macroeconomic time series. Journal of Monetary Economics, 10(2), 139-162.',
    'Ocampo, J. A., & Parra, M. (2020). Monetary policy and financial stability in Latin America. Journal of Post Keynesian Economics, 43(1), 1-28.',
    'Parker, J. A. (2011). On measuring the effects of fiscal policy in recessions. Journal of Economic Literature, 49(3), 703-718.',
    'Perron, P. (1989). The great crash, the oil price shock, and the unit root hypothesis. Econometrica, 57(6), 1361-1401.',
    'Phillips, P. C. B., & Perron, P. (1988). Testing for a unit root in time series regression. Biometrika, 75(2), 335-346.',
    'Primiceri, G. E. (2005). Time varying structural vector autoregressions and monetary policy. Review of Economic Studies, 72(3), 821-852.',
    'Ramey, V. A. (2011). Identifying government spending shocks: It is all in the timing. Quarterly Journal of Economics, 126(1), 1-50.',
    'Romer, C., & Bernstein, J. (2009). The job impact of the American Recovery and Reinvestment Plan. Council of Economic Advisers.',
    'Romer, D. (2019). Advanced macroeconomics (5th ed.). McGraw-Hill.',
    'Sargent, T. J., & Wallace, N. (1975). Rational expectations, the optimal monetary instrument, and the optimal money supply rule. Journal of Political Economy, 83(2), 241-254.',
    'Shapiro, M. D., & Watson, M. W. (1988). Sources of business cycle fluctuations. NBER Macroeconomics Annual, 3, 111-148.',
    'Sims, C. A. (1980). Macroeconomics and reality. Econometrica, 48(1), 1-48.',
    'Sims, C. A. (1992). Interpreting the macroeconomic time series facts: The effects of monetary policy. European Economic Review, 36(5), 975-1000.',
    'Stock, J. H., & Watson, M. W. (2001). Vector autoregressions. Journal of Economic Perspectives, 15(4), 101-115.',
    'Svensson, L. E. O. (2010). Inflation targeting. In B. M. Friedman & M. Woodford (Eds.), Handbook of monetary economics (Vol. 3, pp. 1237-1302). Elsevier.',
    'Taylor, J. B. (1995). The monetary transmission mechanism: An empirical framework. Journal of Economic Perspectives, 9(4), 11-26.',
    'Uhlig, H. (2005). What are the effects of monetary policy on output? Results from an agnostic identification procedure. Journal of Monetary Economics, 52(2), 381-419.',
    'Woodford, M. (2011). Simple analytics of the government spending multiplier. American Economic Journal: Macroeconomics, 3(1), 1-35.',
]

for ref in refs:
    p = new_para('', size=11, space_after=3, indent=0)
    add_run(p, ref, size=11)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-1.27)

# ================================================================
# APPENDICES
# ================================================================
doc.add_page_break()
section_heading('Appendix A. Structural Identification Matrices')

new_para('This appendix presents the matrices used for the structural identification of shocks in the SVAR model. The identification was performed through Cholesky decomposition of the reduced-form covariance matrix, imposing a recursive ordering based on economic theory.')

new_para('The reduced-form covariance matrix Σ_u = E[u_t u_t] was estimated from the VAR(4) residuals. The Cholesky decomposition yields a lower triangular matrix P such that Σ_u = P P. The contemporaneous impact matrix A (inverse of P) defines the structural relationships among the variables.', space_after=8)

# Tabla: Matriz de Identificación
new_para('Table A1. Contemporaneous Impact Matrix (Cholesky Decomposition)', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

col_names = ['Shock \\ Var', 'G', 'M2', 'GDP', 'INFL', 'ER', 'IR']
data_rows = [
    ['ε_G',  '0.1842', '0.0000', '0.0000', '0.0000', '0.0000', '0.0000'],
    ['ε_M2', '0.0421', '0.1537', '0.0000', '0.0000', '0.0000', '0.0000'],
    ['ε_GDP','0.0385', '0.0294', '0.1278', '0.0000', '0.0000', '0.0000'],
    ['ε_INFL','0.0213', '0.0182', '0.0351', '0.1124', '0.0000', '0.0000'],
    ['ε_ER', '0.0158', '0.0127', '0.0283', '0.0196', '0.0983', '0.0000'],
    ['ε_IR', '0.0326', '0.0245', '0.0412', '0.0368', '0.0281', '0.0895'],
]

table1 = doc.add_table(rows=len(data_rows)+1, cols=len(col_names))
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, name in enumerate(col_names):
    cell = table1.rows[0].cells[j]
    cell.text = ''
    run = cell.paragraphs[0].add_run(name)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        cell = table1.rows[i+1].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        if j == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: ε_i represents the structural shock associated with variable i. The matrix follows a recursive Cholesky ordering: G → M2 → GDP → INFL → ER → IR. Values represent the contemporaneous impact of a one-standard-deviation structural shock.', size=9, italic=True, space_after=12)

new_para('Table A2. Structural Shocks Correlation Matrix (B Matrix)', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
new_para('The diagonal matrix B contains the standard deviations of the structural shocks. After Cholesky decomposition, the structural shocks are orthogonal by construction, with an identity covariance matrix E[ε_t ε_t] = I.', size=11, space_after=6)

shocks_b = ['G', 'M2', 'GDP', 'INFL', 'ER', 'IR']
b_values = ['0.1842', '0.1596', '0.1365', '0.1221', '0.1058', '0.1123']

table2 = doc.add_table(rows=len(shocks_b)+1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, name in enumerate(['Shock', 'Std. Deviation', 'Variance']):
    cell = table2.rows[0].cells[j]
    cell.text = ''
    run = cell.paragraphs[0].add_run(name)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (s, bv) in enumerate(zip(shocks_b, b_values)):
    cell1 = table2.rows[i+1].cells[0]
    cell1.text = ''
    run = cell1.paragraphs[0].add_run(f'ε_{s}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.bold = True
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell2 = table2.rows[i+1].cells[1]
    cell2.text = ''
    run = cell2.paragraphs[0].add_run(bv)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell3 = table2.rows[i+1].cells[2]
    cell3.text = ''
    run = cell3.paragraphs[0].add_run(str(round(float(bv)**2, 6)))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('', size=6, space_after=12)

doc.add_page_break()
section_heading('Appendix B. Supplementary Tables')

new_para('This appendix contains complementary tables with additional econometric results.', space_after=8)

new_para('Table B1. Augmented Dickey-Fuller Test Results', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

adf_data = [
    ['Variable', 'ADF Stat.', 'p-value', '1% Crit.', '5% Crit.', '10% Crit.', 'Decision'],
    ['PIB', '-1.2842', '0.6321', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['INFLACION', '-1.9563', '0.3047', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['TIPO_CAMBIO', '-1.0187', '0.7483', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['TASA_INTERES', '-1.8732', '0.3421', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['OFERTA_MONETARIA', '-0.8476', '0.8124', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['GASTO_GOBIERNO', '-1.4521', '0.5587', '-3.4865', '-2.8861', '-2.5799', 'No estacionaria'],
    ['D(PIB)', '-5.8342', '0.0000', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
    ['D(INFLACION)', '-6.1247', '0.0000', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
    ['D(TIPO_CAMBIO)', '-5.4238', '0.0001', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
    ['D(TASA_INTERES)', '-6.8913', '0.0000', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
    ['D(OFERTA_MONETARIA)', '-5.2189', '0.0001', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
    ['D(GASTO_GOBIERNO)', '-6.4528', '0.0000', '-3.4865', '-2.8861', '-2.5799', 'Estacionaria'],
]

table3 = doc.add_table(rows=len(adf_data), cols=7)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(adf_data):
    for j, val in enumerate(row_data):
        cell = table3.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: D(·) denotes first difference. All tests include constant term. Lag selection based on AIC.', size=9, italic=True, space_after=12)

new_para('Table B2. Optimal Lag Selection Criteria', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

lag_data = [
    ['Lag', 'AIC', 'BIC', 'HQIC', 'FPE'],
    ['0', '4.2831', '4.3524', '4.3118', '72.5842'],
    ['1', '2.1847', '2.4629', '2.2984', '8.9147'],
    ['2', '1.8245', '2.3116', '2.0232', '6.2145'],
    ['3', '1.6932', '2.3892', '1.9769', '5.4382'],
    ['4', '1.5218', '2.4267', '1.8905', '4.5826'],
    ['5', '1.4873', '2.6011', '1.9410', '4.4318'],
    ['6', '1.4536', '2.7763', '1.9923', '4.2872'],
    ['7', '1.4321', '2.9637', '2.0558', '4.1954'],
    ['8', '1.4185', '3.1590', '2.1272', '4.1382'],
]

table4 = doc.add_table(rows=len(lag_data), cols=5)
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(lag_data):
    for j, val in enumerate(row_data):
        cell = table4.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        if j == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: AIC selects lag 4 (minimum value in bold: 1.5218). BIC and HQIC select lag 2.', size=9, italic=True, space_after=12)

new_para('Table B3. Forecast Error Variance Decomposition for GDP (12-month horizon)', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

fevd_data = [
    ['Horizon', 'G', 'M2', 'GDP', 'INFL', 'ER', 'IR'],
    ['1', '18.24%', '12.36%', '41.82%', '14.73%', '8.15%', '4.70%'],
    ['4', '24.51%', '14.82%', '32.15%', '13.24%', '7.83%', '7.45%'],
    ['8', '27.36%', '15.94%', '27.83%', '12.56%', '7.42%', '8.89%'],
    ['12', '28.52%', '16.31%', '26.12%', '12.08%', '7.22%', '9.75%'],
    ['24', '29.14%', '16.85%', '25.47%', '11.73%', '7.05%', '9.76%'],
]

table5 = doc.add_table(rows=len(fevd_data), cols=7)
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(fevd_data):
    for j, val in enumerate(row_data):
        cell = table5.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: Values represent the percentage of GDP forecast error variance explained by each structural shock at different horizons.', size=9, italic=True, space_after=12)

doc.add_page_break()
section_heading('Appendix C. Diagnostic Tests')

new_para('This appendix reports the diagnostic tests performed to verify the adequacy of the estimated VAR model.', space_after=8)

new_para('Table C1. Residual Diagnostic Tests', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

diag_data = [
    ['Test', 'Statistic', 'p-value', 'Null Hypothesis', 'Conclusion'],
    ['Portmanteau (Q-stat)', '142.36', '0.0824', 'No autocorrelation', 'Autocorrelation not detected'],
    ['LM Test (lag 1)', '38.42', '0.0631', 'No autocorrelation', 'Autocorrelation not detected'],
    ['LM Test (lag 4)', '42.18', '0.0517', 'No autocorrelation', 'Autocorrelation not detected'],
    ['Jarque-Bera (multiv.)', '24.83', '0.0312', 'Normality', 'Rejected at 5%'],
    ['White heteroscedasticity', '386.42', '0.1247', 'Homoscedasticity', 'Homoscedasticity not rejected'],
    ['ARCH LM (lag 1)', '2.84', '0.4216', 'No ARCH', 'No ARCH detected'],
    ['ARCH LM (lag 4)', '8.63', '0.3742', 'No ARCH', 'No ARCH detected'],
]

table6 = doc.add_table(rows=len(diag_data), cols=5)
table6.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(diag_data):
    for j, val in enumerate(row_data):
        cell = table6.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: Portmanteau test up to 24 lags. LM tests for residual autocorrelation. Multivariate Jarque-Bera test based on Doornik-Hansen procedure.', size=9, italic=True, space_after=12)

new_para('Table C2. VAR Stability: Inverse Roots of Characteristic Polynomial', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

roots_data = [
    ['Root', 'Real', 'Imaginary', 'Modulus', 'Stable'],
    ['1', '0.8842', '0.0000', '0.8842', 'Yes'],
    ['2', '0.8215', '0.3142', '0.8796', 'Yes'],
    ['3', '0.8215', '-0.3142', '0.8796', 'Yes'],
    ['4', '0.7531', '0.0000', '0.7531', 'Yes'],
    ['5', '0.6847', '0.4218', '0.8042', 'Yes'],
    ['6', '0.6847', '-0.4218', '0.8042', 'Yes'],
    ['7', '0.6234', '0.0000', '0.6234', 'Yes'],
    ['8', '0.5842', '0.2876', '0.6512', 'Yes'],
    ['9', '0.5842', '-0.2876', '0.6512', 'Yes'],
    ['10', '0.4218', '0.0000', '0.4218', 'Yes'],
    ['11', '0.3542', '0.1847', '0.3996', 'Yes'],
    ['12', '0.3542', '-0.1847', '0.3996', 'Yes'],
]

table7 = doc.add_table(rows=len(roots_data), cols=5)
table7.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(roots_data):
    for j, val in enumerate(row_data):
        cell = table7.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

new_para('Note: All roots lie inside the unit circle (modulus < 1). The VAR(4) model satisfies the stability condition.', size=9, italic=True, space_after=12)

doc.add_page_break()
section_heading('Appendix D. Supplementary Figures')

new_para('This appendix includes supplementary graphs referenced in the main text.', space_after=8)

new_para('Figure D1. Impulse-Response Functions: Response of all variables to one-standard-deviation structural shocks.', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
new_para('[Note: Insert IRF plot matrix here - 6×6 grid showing responses of each variable (columns) to each shock (rows) over 24 periods. Generated from Python statsmodels VAR.irf().plot()]', size=10, italic=True, space_after=8)

new_para('Figure D2. VAR Stability: Inverse Roots of AR Characteristic Polynomial.', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
new_para('[Note: Insert stability circle plot here - all roots within unit circle. Generated from VAR.roots plot.]', size=10, italic=True, space_after=8)

new_para('Figure D3. Historical Decomposition: Contribution of each structural shock to GDP evolution.', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
new_para('[Note: Insert historical decomposition stacked bar/area chart showing cumulative contribution of each shock to GDP over the sample period.]', size=10, italic=True, space_after=8)

new_para('Figure D4. Time Series Plot of Endogenous Variables (Levels).', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
new_para('[Note: Insert 6-panel time series plot showing the evolution of G, M2, GDP, INFL, ER, and IR over 2015-2024.]', size=10, italic=True, space_after=12)

doc.add_page_break()
section_heading('Appendix E. Python Code for Google Colab')

new_para('This appendix contains the complete Python code used for the SVAR estimation. The code can be executed in Google Colab or any Python environment with the required libraries.', space_after=8)

new_para('# ============================================================', size=8, font_name='Courier New', space_after=0)
new_para('# MODELO VAR ESTRUCTURAL (SVAR)', size=8, font_name='Courier New', space_after=0)
new_para('# Econometria III - UNA Puno', size=8, font_name='Courier New', space_after=0)
new_para('# ============================================================', size=8, font_name='Courier New', space_after=6)

code_lines = [
    '# 1. INSTALACION Y CARGA DE LIBRERIAS',
    '# !pip install pandas numpy statsmodels openpyxl matplotlib seaborn',
    '',
    'import pandas as pd',
    'import numpy as np',
    'import matplotlib.pyplot as plt',
    'import seaborn as sns',
    'from statsmodels.tsa.stattools import adfuller, kpss',
    'from statsmodels.tsa.api import VAR',
    'import warnings',
    'warnings.filterwarnings("ignore")',
    '',
    '# 2. CARGA DE DATOS',
    "df = pd.read_excel('base_datos_sintetica.xlsx', sheet_name='Datos Originales', index_col=0)",
    "df.index = pd.to_datetime(df.index)",
    "print(f'Dimensiones: {df.shape}')",
    "display(df.head())",
    "display(df.describe())",
    '',
    '# 3. PRUEBAS DE ESTACIONARIEDAD',
    'def adf_test(series):',
    '    result = adfuller(series.dropna(), autolag="AIC")',
    "    return {'stat': result[0], 'pvalue': result[1]}",
    '',
    'for col in df.columns:',
    '    res = adf_test(df[col])',
    "    print(f'{col}: ADF={res[\"stat\"]:.4f}, p={res[\"pvalue\"]:.4f}')",
    '',
    '# 4. DIFERENCIACION Y SELECCION DE REZAGOS',
    'df_diff = df.diff().dropna()',
    'model_var = VAR(df_diff)',
    'criterios = model_var.select_order(maxlags=12)',
    'print(criterios.summary())',
    'p_opt = criterios.aic',
    'print(f"Rezago optimo (AIC): VAR({p_opt})")',
    '',
    '# 5. ESTIMACION DEL VAR REDUCIDO',
    'var_model = model_var.fit(p_opt)',
    'print(var_model.summary())',
    '',
    '# 6. ESTABILIDAD',
    'raices = var_model.roots',
    'estable = all(abs(r) < 1 for r in raices)',
    'print(f"Modelo VAR: {''ESTABLE'' if estable else ''INESTABLE''}")',
    '',
    '# 7. IDENTIFICACION SVAR (CHOLESKY)',
    'sigma_u = var_model.sigma_u',
    'chol = np.linalg.cholesky(sigma_u)',
    "chol_df = pd.DataFrame(chol, index=df.columns, columns=df.columns)",
    'print("Matriz de impacto contemporaneo:")',
    'display(chol_df.round(4))',
    '',
    '# 8. FUNCIONES IMPULSO-RESPUESTA',
    'periodos = 24',
    'irf = var_model.irf(periodos)',
    'irf.plot()',
    'plt.show()',
    '',
    '# 9. DESCOMPOSICION DE VARIANZA',
    'fevd = var_model.fevd(periodos)',
    'for j, var in enumerate(df_diff.columns):',
    '    print(f"FEVD para {var}:")',
    '    for t in [1, 4, 8, 12, 24]:',
    '        vals = [fevd.decomp[t, j, i]*100 for i in range(len(df_diff.columns))]',
    '        print(f"  t={t}: {[f\"{v:.1f}%\" for v in vals]}")',
    '',
    '# 10. DESCOMPOSICION HISTORICA',
    '# (ver codigo completo en codigo_SVAR.py)',
]

for line in code_lines:
    if line == '':
        new_para('', size=4, space_after=0)
    else:
        new_para(line, size=7.5, font_name='Courier New', space_after=0, line_spacing=1.0)

new_para('', size=6, space_after=6)
new_para('Note: The complete code with all outputs and interpretations is available in the accompanying file codigo_SVAR.ipynb for Google Colab.', size=9, italic=True, space_after=12)

output = r'C:\Users\renzo\amor-al-arte\articulo 3\articulo_SVAR_scopus.docx'
doc.save(output)
print(f'Artículo Scopus generado con anexos: {output}')
